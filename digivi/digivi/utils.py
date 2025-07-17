import pandas as pd
import numpy as np
import io
import base64
import matplotlib.pyplot as plt
from datetime import datetime
import gspread
from google.oauth2.service_account import Credentials
import pandas as pd
import logging
from django.conf import settings
from io import BytesIO
import plotly.graph_objects as go
import plotly.express as px

logger = logging.getLogger(__name__)

# ANSI color codes for terminal output
class Colors:
    GREEN = '\033[92m'
    RED = '\033[91m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    WHITE = '\033[97m'
    BOLD = '\033[1m'
    END = '\033[0m'

def print_status(message, status="info"):
    """Print colored status messages with icons"""
    if status == "success":
        print(f"{Colors.GREEN}✓ {message}{Colors.END}")
    elif status == "error":
        print(f"{Colors.RED}✗ {message}{Colors.END}")
    elif status == "warning":
        print(f"{Colors.YELLOW}⚠ {message}{Colors.END}")
    elif status == "info":
        print(f"{Colors.BLUE}ℹ {message}{Colors.END}")
    elif status == "process":
        print(f"{Colors.CYAN}⚡ {message}{Colors.END}")

def get_google_sheets_client():
    """
    Create and return a Google Sheets client using service account credentials.
    
    Returns:
        gspread.Client: Authenticated Google Sheets client
    """
    try:
        print_status("Initializing Google Sheets client...", "process")
        
        # Define the scope for Google Sheets API
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/drive"
        ]
        print_status("API scopes configured", "success")
        
        # Load credentials from service account key file
        # Make sure to set GOOGLE_SERVICE_ACCOUNT_KEY_PATH in your settings
        print_status("Loading service account credentials...", "process")
        credentials = Credentials.from_service_account_file(
            settings.GOOGLE_SERVICE_ACCOUNT_KEY_PATH, 
            scopes=scope
        )
        print_status("Service account credentials loaded", "success")
        
        # Create and return the client
        print_status("Authorizing Google Sheets client...", "process")
        client = gspread.authorize(credentials)
        print_status("Google Sheets client authorized successfully", "success")
        return client
    
    except Exception as e:
        print_status(f"Failed to create Google Sheets client: {e}", "error")
        logger.error(f"Error creating Google Sheets client: {e}")
        return None

def get_google_sheets_client_from_dict():
    """
    Alternative method: Create Google Sheets client from credentials dictionary.
    Use this if you store credentials as a dictionary in settings instead of a file.
    
    Returns:
        gspread.Client: Authenticated Google Sheets client
    """
    try:
        print_status("Initializing Google Sheets client from dict...", "process")
        
        # Define the scope for Google Sheets API
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/drive"
        ]
        print_status("API scopes configured", "success")
        
        # Load credentials from dictionary
        # Make sure to set GOOGLE_SERVICE_ACCOUNT_INFO as a dict in your settings
        print_status("Loading credentials from dictionary...", "process")
        credentials = Credentials.from_service_account_info(
            settings.GOOGLE_SERVICE_ACCOUNT_INFO, 
            scopes=scope
        )
        print_status("Credentials loaded from dictionary", "success")
        
        # Create and return the client
        print_status("Authorizing Google Sheets client...", "process")
        client = gspread.authorize(credentials)
        print_status("Google Sheets client authorized successfully", "success")
        return client
    
    except Exception as e:
        print_status(f"Failed to create Google Sheets client from dict: {e}", "error")
        logger.error(f"Error creating Google Sheets client from dict: {e}")
        return None

def load_master_from_google_sheets():
    """
    Load master data from Google Sheets and return as a dictionary of DataFrames.
    
    Returns:
        dict: Dictionary with sheet names as keys and DataFrames as values
        None: If loading fails
    """
    try:
        print_status("Starting master data load from Google Sheets...", "process")
        
        # Get the Google Sheets client
        client = get_google_sheets_client()
        if not client:
            print_status("Google Sheets client creation failed", "error")
            logger.error("Failed to create Google Sheets client")
            return None
        
        # Open the spreadsheet using the ID from settings
        print_status("Opening spreadsheet...", "process")
        spreadsheet = client.open_by_key(settings.MASTER_DATA_SHEET_ID)
        print_status(f"Spreadsheet opened: {spreadsheet.title}", "success")
        
        # Get all worksheets
        print_status("Retrieving worksheets...", "process")
        worksheets = spreadsheet.worksheets()
        print_status(f"Found {len(worksheets)} worksheets", "success")
        
        master_data = {}
        
        for worksheet in worksheets:
            try:
                print_status(f"Processing worksheet: {worksheet.title}", "process")
                
                # Get all values from the worksheet
                values = worksheet.get_all_values()
                
                if not values:
                    print_status(f"Worksheet '{worksheet.title}' is empty", "warning")
                    logger.warning(f"Worksheet '{worksheet.title}' is empty")
                    continue
                
                # Convert to DataFrame
                # First row is assumed to be headers
                headers = values[0]
                data = values[1:] if len(values) > 1 else []
                
                df = pd.DataFrame(data, columns=headers)
                
                # Clean up empty rows and columns
                df = df.dropna(how='all')  # Remove completely empty rows
                df = df.loc[:, ~df.columns.str.contains('^Unnamed')]  # Remove unnamed columns
                
                # Convert numeric columns where possible
                for col in df.columns:
                    df[col] = pd.to_numeric(df[col], errors='ignore')
                
                master_data[worksheet.title] = df
                print_status(f"Worksheet '{worksheet.title}' loaded successfully ({len(df)} rows)", "success")
                logger.info(f"Successfully loaded worksheet '{worksheet.title}' with {len(df)} rows")
                
            except Exception as e:
                print_status(f"Failed to load worksheet '{worksheet.title}': {e}", "error")
                logger.error(f"Error loading worksheet '{worksheet.title}': {e}")
                continue
        
        if master_data:
            print_status(f"Master data load completed - {len(master_data)} worksheets loaded", "success")
            logger.info(f"Successfully loaded {len(master_data)} worksheets from Google Sheets")
            return master_data
        else:
            print_status("No worksheets were successfully loaded", "error")
            logger.error("No worksheets were successfully loaded")
            return None
            
    except Exception as e:
        print_status(f"Master data load failed: {e}", "error")
        logger.error(f"Error loading master data from Google Sheets: {e}")
        return None

def refresh_master_data_cache(request):
    """
    Refresh the master data cache by loading fresh data from Google Sheets.
    
    Args:
        request: Django request object (for session access)
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        print_status("Starting master data cache refresh...", "process")
        
        master25 = load_master_from_google_sheets()
        if master25:
            print_status("Converting data to JSON for session storage...", "process")
            request.session['master25'] = {
                name: df.to_json(date_format='iso') for name, df in master25.items()
            }
            print_status("Master data cache refreshed successfully", "success")
            logger.info("Master data cache refreshed successfully")
            return True
        else:
            print_status("Master data cache refresh failed", "error")
            logger.error("Failed to refresh master data cache")
            return False
    except Exception as e:
        print_status(f"Cache refresh error: {e}", "error")
        logger.error(f"Error refreshing master data cache: {e}")
        return False

def get_sheet_info(sheet_id):
    """
    Get basic information about a Google Sheet.
    
    Args:
        sheet_id (str): Google Sheets ID
    
    Returns:
        dict: Sheet information including title, worksheets, etc.
        None: If operation fails
    """
    try:
        print_status("Retrieving sheet information...", "process")
        
        client = get_google_sheets_client()
        if not client:
            print_status("Failed to get Google Sheets client", "error")
            return None
        
        print_status("Opening spreadsheet for info retrieval...", "process")
        spreadsheet = client.open_by_key(sheet_id)
        print_status(f"Spreadsheet accessed: {spreadsheet.title}", "success")
        
        info = {
            'title': spreadsheet.title,
            'id': spreadsheet.id,
            'url': spreadsheet.url,
            'worksheets': []
        }
        
        print_status("Collecting worksheet information...", "process")
        for worksheet in spreadsheet.worksheets():
            worksheet_info = {
                'title': worksheet.title,
                'id': worksheet.id,
                'row_count': worksheet.row_count,
                'col_count': worksheet.col_count
            }
            info['worksheets'].append(worksheet_info)
            print_status(f"Worksheet info collected: {worksheet.title} ({worksheet.row_count}x{worksheet.col_count})", "success")
        
        print_status("Sheet information retrieval completed", "success")
        return info
    
    except Exception as e:
        print_status(f"Failed to get sheet info: {e}", "error")
        logger.error(f"Error getting sheet info: {e}")
        return None

def validate_google_sheets_connection():
    """
    Validate that Google Sheets API connection is working.
    
    Returns:
        tuple: (success: bool, message: str)
    """
    try:
        print_status("Validating Google Sheets connection...", "process")
        
        client = get_google_sheets_client()
        if not client:
            print_status("Google Sheets client validation failed", "error")
            return False, "Failed to create Google Sheets client"
        
        # Try to access the master sheet
        print_status("Testing access to master sheet...", "process")
        spreadsheet = client.open_by_key(settings.MASTER_DATA_SHEET_ID)
        
        success_msg = f"Successfully connected to sheet: {spreadsheet.title}"
        print_status(success_msg, "success")
        return True, success_msg
    
    except Exception as e:
        error_msg = f"Connection failed: {str(e)}"
        print_status(error_msg, "error")
        return False, error_msg

def error_fix(df):
    df = df.copy() 
    to_drop = set()
    col = df.columns[1]

    i = 0
    while i <= len(df) - 3:
        idx_i = df.index[i]
        idx_ip1 = df.index[i+1]
        idx_ip2 = df.index[i+2]

        if idx_i in to_drop or idx_ip1 in to_drop or idx_ip2 in to_drop:
            i += 1
            continue

        v_i = df.loc[idx_i, col]
        v_ip1 = df.loc[idx_ip1, col]
        v_ip2 = df.loc[idx_ip2, col]

        if v_i > v_ip1 and v_i > v_ip2:
            to_drop.add(idx_i)
        elif v_i > v_ip1 and v_i < v_ip2:
            to_drop.add(idx_ip1)
        i += 1

    # Additional rule: second last > last
    if len(df) >= 2:
        second_last_idx = df.index[-2]
        last_idx = df.index[-1]
        if df.loc[second_last_idx, col] > df.loc[last_idx, col]:
            to_drop.add(second_last_idx)

    return df.drop(index=to_drop).reset_index(drop=True)

def get_filled(meter_name, filtered_data, date_column, start_date, end_date, plot_size):
    if not meter_name:
        raise ValueError("Meter name missing.")
        
    # Filter data for the specific meter
    df = filtered_data[filtered_data['Meter Serial Number - as shown on meter'] == meter_name].copy()
    if df.empty:
        return pd.DataFrame()
        
    # Keep only relevant columns
    df = df[[date_column, 'Reading in the meter - in m3']].dropna()

    # Convert date column to datetime
    df[date_column] = pd.to_datetime(df[date_column])

    # Sort by date (if timestamps existed, include them in sort)
    df = df.sort_values(date_column)

    # Deduplicate: keep the latest reading if multiple on same date
    df = df.drop_duplicates(subset=[date_column], keep='last')

    df = error_fix(df)

    # Calculate time difference and deltas
    df['Days Since Previous Reading'] = df[date_column].diff().dt.days.fillna(method='bfill').fillna(1).astype(int)
    df['Delta m³'] = df['Reading in the meter - in m3'].diff().fillna(method='bfill').fillna(0)

    if plot_size <= 0:
        raise ValueError("Invalid plot size.")

    # Normalize per acre and per day
    df['m³ per Acre'] = df['Delta m³'] / plot_size
    df['m³ per Acre per Avg Day'] = df['m³ per Acre'] / df['Days Since Previous Reading'].replace(0, 1)


    # Create filled daily dataframe
    filled = create_adjusted_filled_dataframe(
        meter_data=df,
        start_date=start_date,
        end_date=end_date,
        date_column=date_column,
        avg_day_column='m³ per Acre per Avg Day'
    )

    return filled

def get_dates(meter_name, filtered_data, date_column, start_date, end_date, plot_size):
    if not meter_name:
        raise ValueError("Meter name missing.")
        
    # Filter data for the specific meter
    df = filtered_data[filtered_data['Meter Serial Number - as shown on meter'] == meter_name].copy()
    if df.empty:
        return pd.DataFrame()
        
    # Keep only relevant columns
    df = df[[date_column, 'Reading in the meter - in m3']].dropna()

    # Convert date column to datetime
    df[date_column] = pd.to_datetime(df[date_column])

    # Sort by date (if timestamps existed, include them in sort)
    df = df.sort_values(date_column)

    # Deduplicate: keep the latest reading if multiple on same date
    df = df.drop_duplicates(subset=[date_column], keep='last')

    df = error_fix(df)

    # Calculate time difference and deltas
    df['Days Since Previous Reading'] = df[date_column].diff().dt.days.fillna(method='bfill').fillna(1).astype(int)
    df['Delta m³'] = df['Reading in the meter - in m3'].diff().fillna(method='bfill').fillna(0)

    if plot_size <= 0:
        raise ValueError("Invalid plot size.")

    # Normalize per acre and per day
    df['m³ per Acre'] = df['Delta m³'] / plot_size
    df['m³ per Acre per Avg Day'] = df['m³ per Acre'] / df['Days Since Previous Reading'].replace(0, 1)

    filtered = df[(df[date_column] >= start_date) & (df[date_column] <= end_date)]

    return filtered



def create_adjusted_filled_dataframe(meter_data, start_date, end_date, date_column, avg_day_column):
    # shift first reading to align with start_date
    filtered = meter_data[meter_data[date_column] >= start_date]
    
    # Get the last date in the filtered DataFrame
    last_date_in_df = filtered['Date'].max()

    # Update end_date if needed
    if last_date_in_df < end_date:
        end_date = last_date_in_df


    if filtered.empty:
        return filtered

    if not filtered.empty:
        first_index = filtered.index[0]
        days_diff = (filtered.iloc[0][date_column] - start_date).days + 1
            
        # Get column names explicitly
        col_m3_per_acre = 'm³ per Acre'
        col_avg_day = 'm³ per Acre per Avg Day'

        # Update values explicitly using column names
        meter_data.loc[first_index, col_m3_per_acre] = meter_data.loc[first_index, col_m3_per_acre]
        meter_data.loc[first_index, col_avg_day] = meter_data.loc[first_index, col_avg_day] / days_diff

        
    all_dates = pd.date_range(start=start_date, end=end_date, freq='D')
    base = pd.DataFrame({date_column: all_dates})
    merged = pd.merge(base, meter_data, on=date_column, how='left')

    last_date = start_date
    for idx, row in merged.iterrows():
        if pd.isna(row[avg_day_column]):
            next_rows = meter_data[meter_data[date_column] > row[date_column]]
            if not next_rows.empty:
                merged.iloc[idx, 1:] = next_rows.iloc[0, 1:]
            else:
                merged.iloc[idx, 1:] = 0
        else:
            last_date = row[date_column]
    
    return merged

def kharif2024_farms(master_df):
    meters_df = master_df['Meters'][['Kharif 24 Meter Serial No', 'Kharif 24 FARMID']].copy()
    meters_df = meters_df.dropna(subset=['Kharif 24 Meter Serial No', 'Kharif 24 FARMID'])
    meters_df = meters_df[
        (meters_df['Kharif 24 Meter Serial No'].astype(str).str.strip() != '') &
        (meters_df['Kharif 24 FARMID'].astype(str).str.strip() != '')
    ]
    farm_to_meters = meters_df.groupby('Kharif 24 FARMID')['Kharif 24 Meter Serial No'].apply(list)
    return farm_to_meters.to_dict()


def encode_plot_to_base64(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png')
    buf.seek(0)
    return base64.b64encode(buf.getvalue()).decode('utf-8')


def get_2024plots(selected_farm, meter_df, master_df, meter_list):
    kharif_df = master_df.get('Kharif 24')
    if kharif_df is None:
        raise ValueError("Kharif 24 sheet not found in master data.")

    farm_row = kharif_df[kharif_df['Kharif 24 FarmID'] == selected_farm]
    if farm_row.empty:
        raise ValueError(f"Farm ID {selected_farm} missing in Kharif 24 data.")

    # determine start/end dates
    if pd.notna(farm_row['Kharif 24 Paddy transplanting date (TPR)'].values[0]):
        start_date = pd.to_datetime(
            farm_row['Kharif 24 Paddy transplanting date (TPR)'].values[0],
            dayfirst=True
        )
    elif pd.notna(farm_row['Kharif 24 Paddy sowing (DSR)'].values[0]):
        start_date = pd.to_datetime(
            farm_row['Kharif 24 Paddy sowing (DSR)'].values[0],
            dayfirst=True
        )
    else:
        start_date = pd.to_datetime('15/06/2024', dayfirst=True)

    if pd.notna(farm_row['Kharif 24 Paddy Harvest date'].values[0]):
        end_date = pd.to_datetime(
            farm_row['Kharif 24 Paddy Harvest date'].values[0],
            dayfirst=True
        )
    else:
        end_date = pd.to_datetime('31/10/2024', dayfirst=True)

    # clean and filter meter data
    raw = meter_df.dropna(axis=1, how='all')
    date_col = raw.columns[0]
    raw[date_col] = pd.to_datetime(raw[date_col], errors='coerce').dt.date
    reading_cols = [c for c in raw.columns if 'reading' in c.lower()]
    filtered = raw[[date_col] + reading_cols].dropna(subset=reading_cols, how='all')

    # get filled dataframes per meter
    plot_size = farm_row['Kharif 24 Acres farm/plot'].values[0]
    filled = {
        m: get_filled_2024(m, filtered, date_col, start_date, end_date, plot_size)
        for m in meter_list
    }

    plots = []
    for m, df in filled.items():
        df['Day'] = (pd.to_datetime(df[date_col]) - start_date).dt.days

        # m³ per Acre
        fig1, ax1 = plt.subplots()
        ax1.plot(df['Day'], df['m³ per Acre'])
        ax1.set(xlabel='Days from start', ylabel='m³ per Acre',
                title=f'm³ per Acre | Meter {m} | Farm {selected_farm}')
        fig1.tight_layout()
        plots.append(fig1)

        # m³ per Acre per Avg Day
        fig2, ax2 = plt.subplots()
        ax2.plot(df['Day'], df['m³ per Acre per Avg Day'])
        ax2.set(xlabel='Days from start', ylabel='m³ per Acre per Avg Day',
                title=f'm³ per Acre per Avg Day | Meter {m} | Farm {selected_farm}')
        fig2.tight_layout()
        plots.append(fig2)

    return plots


def create_adjusted_filled_dataframe_2024(meter_data, start_date, end_date, date_column, avg_day_column):
    # shift first reading to align with start_date
    filtered = meter_data[meter_data[date_column] > start_date]
    if not filtered.empty:
        days_diff = (filtered.iloc[0, 0] - start_date).days + 1

        meter_data.loc[
            filtered.index[0],
            meter_data.columns[-2:]
        ] = [
            meter_data.iloc[filtered.index[0], -2],
            meter_data.iloc[filtered.index[0], -1] / days_diff
        ]

    all_dates = pd.date_range(start=start_date, end=end_date, freq='D')
    base = pd.DataFrame({date_column: all_dates})
    merged = pd.merge(base, meter_data, on=date_column, how='left')

    last_date = start_date
    for idx, row in merged.iterrows():
        if pd.isna(row[avg_day_column]):
            next_rows = meter_data[meter_data[date_column] > row[date_column]]
            if not next_rows.empty:
                merged.iloc[idx, 2:] = next_rows.iloc[0, 2:]
            else:
                merged.iloc[idx, 2:] = 0
        else:
            last_date = row[date_column]

    return merged


def get_filled_2024(meter_name, filtered_data, date_column, start_date, end_date, plot_size):
    if not meter_name:
        raise ValueError("Meter name missing.")

    target = f"{meter_name} - reading"
    cols = [c for c in filtered_data.columns if c == date_column or target in c]
    df = filtered_data[cols].dropna(subset=[cols[-1]]).sort_values(date_column)
    df[date_column] = pd.to_datetime(df[date_column])
    df['Days Since Previous Reading'] = df[date_column].diff().dt.days.fillna(method='bfill').astype(int)
    df['Delta m³'] = df[cols[-1]].diff().fillna(method='bfill')

    if plot_size <= 0:
        raise ValueError("Invalid plot size.")

    df['m³ per Acre'] = df['Delta m³'] / plot_size
    df['m³ per Acre per Avg Day'] = df['m³ per Acre'] / df['Days Since Previous Reading'].replace(0, 1)

    return create_adjusted_filled_dataframe_2024(
        meter_data=df,
        start_date=start_date,
        end_date=end_date,
        date_column=date_column,
        avg_day_column='m³ per Acre per Avg Day'
    )


# --- 2025 Helpers ---

def kharif2025_farms(master_df):
    sheet = master_df.get('Farm details')
    if sheet is None:
        raise ValueError("‘Kharif 25’ sheet not found.")
    mapping = {}
    for _, row in sheet.iterrows():
        farm = row.get('Kharif 25 Farm ID')

        if pd.isna(farm):
            continue
        meters = []
        for col in ['Kharif 25 Meter serial number - 1', 'Kharif 25 Meter serial number - 2']:
            serial = row.get(col)
            if pd.notna(serial):
                meters.append(str(serial).strip())
        if meters:
            mapping[str(farm)] = meters

    return mapping


def get_2025plots(raw_df, master_df, selected_farm, meter_list, start_date_enter = None, end_date_enter = None):
    meta = master_df['Farm details']
    farm_row = meta[meta['Kharif 25 Farm ID'] == selected_farm]
    if farm_row.empty:
        raise ValueError(f"Farm ID {selected_farm} not found in metadata.")
    farm_row = farm_row.iloc[0]

    acreage = farm_row.get('Kharif 25 Acres farm - farmer reporting') or 1
    if pd.isna(acreage) or acreage <= 0:
        acreage = 1
        print("acreage is nan")
    if pd.notna(farm_row.get('Kharif 25 Paddy transplanting date (TPR)')):
        start_date = pd.to_datetime(farm_row['Kharif 25 Paddy transplanting date (TPR)'], dayfirst=True)
    else:
        start_date = pd.to_datetime('20/06/2025', dayfirst=True)

    if start_date_enter is not None:
        start_date = start_date_enter

    # Clean meter data
    date_col = 'Date'
    raw_df[date_col] = pd.to_datetime(raw_df[date_col], dayfirst=False)
    raw_df = raw_df.sort_values(date_col)

    plots_base64 = []

    for meter in meter_list:
        end_date = pd.to_datetime(datetime.now().date())
        if end_date_enter is not None:
            end_date = end_date_enter
        
        filled_df = get_filled(meter, raw_df, date_col, start_date, end_date, acreage)
        if filled_df.empty:
            continue

        filled_df['Day'] = (pd.to_datetime(filled_df[date_col]) - start_date).dt.days
        filled_df['7-day SMA'] = filled_df['m³ per Acre per Avg Day'].rolling(window=7, min_periods=1).mean()

        # List to collect plots for this meter
        meter_plots = []

        # --- Graph 1: m3 per acre per avg day ---
        fig1, ax1 = plt.subplots(figsize=(14, 8))
        ax1.plot(filled_df['Day'], filled_df['m³ per Acre per Avg Day'], label='Daily m³ per Acre per day', color='blue')
        ax1.set(title=f'Daily Avg per Acre | Meter {meter}', xlabel='Days from transplanting', ylabel='Daily Avg m³ per Acre')
        ax1.legend()
        # Force axis to start from zero
        ax1.set_xlim(left=0)
        ax1.set_ylim(bottom=0)
        fig1.tight_layout()
        meter_plots.append(encode_plot_to_base64(fig1))
        plt.close(fig1)

        # --- Graph 2: Moving Averages Comparison ---
        fig2, ax2 = plt.subplots(figsize=(14, 8))
        ax2.plot(filled_df['Day'], filled_df['7-day SMA'], label='7-day SMA', linestyle='--', color='green')
        # ax2.plot(filled_df['Day'], filled_df['Weekly Avg'], label='Weekly Avg', linestyle=':', color='orange')
        ax2.set(title=f'Moving Average | Meter {meter}', xlabel='Days from transplanting', ylabel='7-days SMA m³ per Acre')
        ax2.legend()
        # Force axis to start from zero
        ax2.set_xlim(left=0)
        ax2.set_ylim(bottom=0)
        fig2.tight_layout()
        meter_plots.append(encode_plot_to_base64(fig2))
        plt.close(fig2)

        # --- Graph 3: Delta Analysis ---
        fig3, ax3 = plt.subplots(figsize=(14, 8))
        ax3.plot(filled_df['Day'], filled_df['Delta m³'], marker='o', linestyle='-', color='purple', label='Delta m³')
        ax3.set(title=f'Meter Actual readings of meter | Meter {meter}', xlabel='Days from transplanting', ylabel='Delta of readings (m³)')
        ax3.legend()
        # Force axis to start from zero
        ax3.set_xlim(left=0)
        ax3.set_ylim(bottom=0)
        fig3.tight_layout()
        meter_plots.append(encode_plot_to_base64(fig3))
        plt.close(fig3)

        # --- Graph 4: Delta per Acre ---
        fig4, ax4 = plt.subplots(figsize=(14, 8))
        ax4.plot(filled_df['Day'], filled_df['m³ per Acre'], marker='x', linestyle='-', color='red', label='Delta m³/Acre')
        ax4.set(title=f'Meter readings per Acre | Meter {meter}', xlabel='Days from transplanting', ylabel='Delta of reading per Acre (m³) ')
        ax4.legend()
        # Force axis to start from zero
        ax4.set_xlim(left=0)
        ax4.set_ylim(bottom=0)
        fig4.tight_layout()
        meter_plots.append(encode_plot_to_base64(fig4))
        plt.close(fig4)

        plots_base64.extend(meter_plots)

    return plots_base64

def get_2025plots_combined(raw_df, master_df, selected_farm, meter_list, start_date_enter=None, end_date_enter=None):
    """
    Generate combined plots for multiple meters treating them as a single unit
    """
    import pandas as pd
    import matplotlib.pyplot as plt
    import base64
    import io
    
    if len(meter_list) < 2:
        return []  # No combined plots for single meter
    
    def encode_plot_to_base64(fig):
        buf = io.BytesIO()
        fig.savefig(buf, format='png')
        buf.seek(0)
        return base64.b64encode(buf.getvalue()).decode('utf-8')
    
    # Get farm metadata
    meta = master_df['Farm details']
    farm_row = meta[meta['Kharif 25 Farm ID'] == selected_farm]
    if farm_row.empty:
        raise ValueError(f"Farm ID {selected_farm} not found in metadata.")
    farm_row = farm_row.iloc[0]
    
    acreage = farm_row.get('Kharif 25 Acres farm - farmer reporting') or 1
    if pd.isna(acreage) or acreage <= 0:
        acreage = 1
    
    if pd.notna(farm_row.get('Kharif 25 Paddy transplanting date (TPR)')):
        start_date = pd.to_datetime(farm_row['Kharif 25 Paddy transplanting date (TPR)'], dayfirst=True)
    else:
        start_date = pd.to_datetime('20/06/2025', dayfirst=True)
    
    if start_date_enter is not None:
        start_date = pd.to_datetime(start_date_enter)
    
    end_date = pd.to_datetime(datetime.now().date())
    if end_date_enter is not None:
        end_date = pd.to_datetime(end_date_enter)
    
    # Clean meter data
    date_col = 'Date'
    raw_df[date_col] = pd.to_datetime(raw_df[date_col], dayfirst=False)
    raw_df = raw_df.sort_values(date_col)
    
    # Get data for all meters
    all_meter_data = []
    for meter in meter_list:
        meter_df = raw_df[raw_df['Meter Serial Number - as shown on meter'] == meter].copy()
        if not meter_df.empty:
            meter_df = meter_df[[date_col, 'Reading in the meter - in m3']].dropna()
            meter_df[date_col] = pd.to_datetime(meter_df[date_col])
            meter_df = meter_df.sort_values(date_col)
            meter_df = meter_df.drop_duplicates(subset=[date_col], keep='last')
            meter_df['Meter'] = meter
            all_meter_data.append(meter_df)
    
    if not all_meter_data:
        return []
    
    # Combine all meter data
    combined_df = pd.concat(all_meter_data, ignore_index=True)
    
    # Group by date and sum readings
    combined_grouped = combined_df.groupby(date_col).agg({
        'Reading in the meter - in m3': 'sum'
    }).reset_index()
    
    # Calculate deltas for combined data
    combined_grouped = combined_grouped.sort_values(date_col)
    combined_grouped['Days Since Previous Reading'] = combined_grouped[date_col].diff().dt.days.fillna(1).astype(int)
    combined_grouped['Delta m³'] = combined_grouped['Reading in the meter - in m3'].diff().fillna(0)
    
    # Normalize per acre
    combined_grouped['m³ per Acre'] = combined_grouped['Delta m³'] / acreage
    combined_grouped['m³ per Acre per Avg Day'] = combined_grouped['m³ per Acre'] / combined_grouped['Days Since Previous Reading'].replace(0, 1)
    
    # Fill daily data
    all_dates = pd.date_range(start=start_date, end=end_date, freq='D')
    base = pd.DataFrame({date_col: all_dates})
    filled_df = pd.merge(base, combined_grouped, on=date_col, how='left')
    
    # Forward fill values
    for col in ['m³ per Acre', 'm³ per Acre per Avg Day', 'Delta m³']:
        filled_df[col] = filled_df[col].fillna(0)
    
    # Calculate days from start
    filled_df['Day'] = (filled_df[date_col] - start_date).dt.days
    filled_df['7-day SMA'] = filled_df['m³ per Acre per Avg Day'].rolling(window=7, min_periods=1).mean()
    
    # Generate plots
    plots_base64 = []
    meters_label = " + ".join(meter_list)
    
    # Graph 1: Daily avg per acre
    fig1, ax1 = plt.subplots(figsize=(14, 8))
    ax1.plot(filled_df['Day'], filled_df['m³ per Acre per Avg Day'], label='Combined Daily m³ per Acre per day', color='blue', linewidth=2)
    ax1.set(title=f'Combined Daily Avg per Acre | Meters: {meters_label}', 
            xlabel='Days from transplanting', 
            ylabel='Daily Avg m³ per Acre')
    ax1.legend()
    ax1.grid(True, alpha=0.3)
    fig1.tight_layout()
    plots_base64.append(encode_plot_to_base64(fig1))
    plt.close(fig1)
    
    # Graph 2: 7-day SMA
    fig2, ax2 = plt.subplots(figsize=(14, 8))
    ax2.plot(filled_df['Day'], filled_df['7-day SMA'], label='Combined 7-day SMA', linestyle='--', color='green', linewidth=2)
    ax2.set(title=f'Combined Moving Average | Meters: {meters_label}', 
            xlabel='Days from transplanting', 
            ylabel='7-days SMA m³ per Acre')
    ax2.legend()
    ax2.grid(True, alpha=0.3)
    fig2.tight_layout()
    plots_base64.append(encode_plot_to_base64(fig2))
    plt.close(fig2)
    
    # Graph 3: Delta Analysis
    fig3, ax3 = plt.subplots(figsize=(14, 8))
    ax3.plot(filled_df['Day'], filled_df['Delta m³'], marker='o', linestyle='-', color='purple', label='Combined Delta m³', linewidth=2)
    ax3.set(title=f'Combined Meter Readings | Meters: {meters_label}', 
            xlabel='Days from transplanting', 
            ylabel='Combined Delta of readings (m³)')
    ax3.legend()
    ax3.grid(True, alpha=0.3)
    fig3.tight_layout()
    plots_base64.append(encode_plot_to_base64(fig3))
    plt.close(fig3)
    
    # Graph 4: Delta per Acre
    fig4, ax4 = plt.subplots(figsize=(14, 8))
    ax4.plot(filled_df['Day'], filled_df['m³ per Acre'], marker='x', linestyle='-', color='red', label='Combined Delta m³/Acre', linewidth=2)
    ax4.set(title=f'Combined Readings per Acre | Meters: {meters_label}', 
            xlabel='Days from transplanting', 
            ylabel='Combined Delta per Acre (m³)')
    ax4.legend()
    ax4.grid(True, alpha=0.3)
    fig4.tight_layout()
    plots_base64.append(encode_plot_to_base64(fig4))
    plt.close(fig4)
    
    return plots_base64

def get_tables(raw_df, master_df, farm_list, col_to_get, start_date_enter = None, end_date_enter = None):
    # Clean meter data
    date_col = 'Date'
    raw_df[date_col] = pd.to_datetime(raw_df[date_col], dayfirst=False)
    raw_df = raw_df.sort_values(date_col)
    dfs = []
    
    for farm in farm_list.keys():
        # Clean meter data
        date_col = 'Date'
        raw_df[date_col] = pd.to_datetime(raw_df[date_col], dayfirst=False)
        raw_df = raw_df.sort_values(date_col)

        meta = master_df['Farm details']
        farm_row = meta[meta['Kharif 25 Farm ID'] == farm]
        if farm_row.empty:
            raise ValueError(f"Farm ID {farm} not found in metadata.")
        farm_row = farm_row.iloc[0]

        acreage = farm_row.get('Kharif 25 Acres farm - farmer reporting') or 1
        if pd.isna(acreage) or acreage <= 0:
            acreage = 1
            print("acreage is nan")
        if pd.notna(farm_row.get('Kharif 25 Paddy transplanting date (TPR)')):
            start_date = pd.to_datetime(farm_row['Kharif 25 Paddy transplanting date (TPR)'], dayfirst=True)
        else:
            start_date = pd.to_datetime('20/06/2025', dayfirst=True)
        
        if start_date_enter is not None:
            start_date = start_date_enter

        
        
        for meter in farm_list[farm]:
            end_date = pd.to_datetime(datetime.now().date())
            if end_date_enter is not None:
                end_date = end_date_enter
            filled_df = get_filled(meter, raw_df, date_col, start_date, end_date, acreage)
            if filled_df.empty:
                filled_df = pd.DataFrame(columns=['Day', meter])
            else:
                filled_df['Day'] = (pd.to_datetime(filled_df[date_col]) - start_date).dt.days
                filled_df = filled_df[['Day', col_to_get]].rename(columns={col_to_get: meter})
            
            dfs.append(filled_df)

    from functools import reduce

    # Assuming all DataFrames have the same name for the date column
    combined_df = reduce(lambda left, right: pd.merge(left, right, on='Day', how='outer'), dfs)

    return combined_df


def calculate_avg_m3_per_acre(group_type, group_label, farm_ids, raw_df, master25, column_to_see, start_date_enter = None, end_date_enter = None):
    """
    Given a group label (like 'Group-A Complied') and its farm IDs, 
    returns a dataframe with Days column and average m³ per acre per day.
    """
    kharif_sheet = master25.get("Farm details")
    all_meter_dfs = []


    # --- Start of main function logic ---

    for farm_id in farm_ids:
        # Get acreage
        acreage_row = kharif_sheet[kharif_sheet["Kharif 25 Farm ID"] == farm_id]
        if acreage_row.empty:
            continue
        acreage = acreage_row["Kharif 25 Acres farm - farmer reporting"].values[0]
        if pd.isna(acreage) or acreage == 0:
            acreage = 1
        # Get transplanting date
        tpr_col = "Kharif 25 Paddy transplanting date (TPR)"
        tpr_date = acreage_row[tpr_col].values[0]
        if pd.isna(tpr_date) or tpr_date == "":
            tpr_date = pd.to_datetime("2025-07-02")
        else:
            tpr_date = pd.to_datetime(tpr_date)
        
        if start_date_enter is not None:
            tpr_date = start_date_enter

        # Get meter serial numbers
        meters = []
        for m_col in ["Kharif 25 Meter serial number - 1", "Kharif 25 Meter serial number - 2"]:
            val = acreage_row[m_col].values[0]
            if pd.notna(val) and val != "":
                meters.append(str(val).strip())

        # Clean meter data
        date_col = 'Date'
        raw_df[date_col] = pd.to_datetime(raw_df[date_col], dayfirst=False)
        raw_df = raw_df.sort_values(date_col)


        # Process each meter
        for meter in meters:
            end_date = pd.to_datetime(datetime.now().date())
          
            if end_date_enter is not None:
                end_date = end_date_enter
            filled_df = get_filled(meter, raw_df, date_col, tpr_date, end_date, acreage)
            if filled_df.empty:
                continue
            filled_df['Day'] = (pd.to_datetime(filled_df[date_col]) - tpr_date).dt.days
            meter_df = filled_df[["Day", column_to_see]].reset_index(drop=True).rename(columns={column_to_see: meter})
            all_meter_dfs.append(meter_df)

    if not all_meter_dfs:
        return pd.DataFrame(columns=["Day", group_label])
    

    # Merge on Days
    merged = pd.DataFrame()
    from functools import reduce

    # Assuming all DataFrames have the same name for the date column
    merged = reduce(lambda left, right: pd.merge(left, right, on='Day', how='outer'), all_meter_dfs)

    # Average across all meters day-wise
    avg_series = merged.drop(columns=["Day"]).mean(axis=1)
    final_df = pd.DataFrame({
        "Day": merged["Day"],
        group_label: avg_series
    })

    return final_df

def generate_group_analysis_plot(df, col_name):
    """
    Takes a DataFrame where:
    - The first column is 'Day'
    - All other columns are group names, with average m³/acre/day values
    Returns base64-encoded image string
    """
    plt.figure(figsize=(16, 10))
    
    # Define distinct markers and colors for better visibility of overlapping points
    markers = ['o', 's', '^', 'D', 'v', '<', '>', 'p', '*', 'h', 'H', '+', 'x']
    colors = plt.cm.Set1(np.linspace(0, 1, len(df.columns[1:])))
    
    for i, col in enumerate(df.columns[1:]):
        plt.plot(df['Day'], df[col], 
                label=col, 
                linewidth=2.5,
                marker=markers[i % len(markers)],
                markersize=8,
                markerfacecolor=colors[i],
                markeredgecolor='black',
                markeredgewidth=1,
                alpha=0.8)
    
    plt.xlabel("Days from Transplanting (TPR)", fontsize=12, fontweight='bold')
    plt.ylabel(col_name, fontsize=12, fontweight='bold')
    plt.title("Group-wise Water Usage Comparison", fontsize=14, fontweight='bold')
    plt.grid(True, alpha=0.3)
    plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
    
    # Force axis to start from zero using current axes
    ax = plt.gca()
    ax.set_xlim(left=0)
    ax.set_ylim(bottom=0)
    plt.tight_layout()
    
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    plt.close()
    encoded = base64.b64encode(buf.read()).decode('utf-8')
    return encoded


        

        
def get_meters_by_village(raw_df, village_name):
    """
    Get all unique meter serial numbers for a specific village.
    Village name is in column D (index 3) of the raw data.
    """
    # Filter rows where village matches
    village_meters = raw_df[raw_df.iloc[:, 3] == village_name]['Meter Serial Number - as shown on meter'].unique()
    return [str(m).strip() for m in village_meters if pd.notna(m)]


def generate_word_report(results, filter_type, filter_value, raw_df, master_df):
    """
    Generate a Word document report with graphs and details
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import base64
    from datetime import datetime
    
    # Create document
    doc = Document()
    
    # Set up styles
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    
    # Heading style
    heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 139)
    
    # Add title
    title = doc.add_paragraph('DIGI-VI Water Meter Analysis Report', style='CustomTitle')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Empty line
    
    # Add filter information
    filter_heading = doc.add_paragraph('Analysis Filter Information', style='CustomHeading')
    
    filter_info = doc.add_paragraph()
    filter_info.add_run('Filter Type: ').bold = True
    filter_info.add_run(f'{filter_type}\n')
    filter_info.add_run('Filter Value: ').bold = True
    filter_info.add_run(f'{filter_value}\n')
    
    if filter_type == "Village Filter":
        # Count total meters in village
        village_meters = raw_df[raw_df.iloc[:, 3] == filter_value]['Meter Serial Number - as shown on meter'].nunique()
        filter_info.add_run('Total Meters in Village: ').bold = True
        filter_info.add_run(f'{village_meters}\n')
    
    doc.add_page_break()
    
    # Process each meter's results
    for idx, block in enumerate(results):
        # Add meter heading
        meter_heading = doc.add_paragraph(f'Meter Unit: {block["meter"]}', style='CustomHeading')
        
        # Add meter details
        details = doc.add_paragraph()
        if 'farm' in block and block['farm']:
            details.add_run('Associated Farm ID: ').bold = True
            details.add_run(f'{block["farm"]}\n')
            
            # Get additional farm details from master data
            farm_details = master_df['Farm details']
            farm_row = farm_details[farm_details['Kharif 25 Farm ID'] == block['farm']]
            if not farm_row.empty:
                farm_row = farm_row.iloc[0]
                
                # Add village name
                village_col = 'Kharif 25 Village'
                if village_col in farm_row and pd.notna(farm_row[village_col]):
                    details.add_run('Village: ').bold = True
                    details.add_run(f'{farm_row[village_col]}\n')
                
                # Add acreage
                acre_col = 'Kharif 25 Acres farm - farmer reporting'
                if acre_col in farm_row and pd.notna(farm_row[acre_col]):
                    details.add_run('Farm Size: ').bold = True
                    details.add_run(f'{farm_row[acre_col]} acres\n')
                
                # Add transplanting date
                tpr_col = 'Kharif 25 Paddy transplanting date (TPR)'
                if tpr_col in farm_row and pd.notna(farm_row[tpr_col]):
                    details.add_run('Transplanting Date: ').bold = True
                    details.add_run(f'{farm_row[tpr_col]}\n')
        
        doc.add_paragraph()  # Empty line
        
        # Add graphs
        graph_titles = [
            "Daily Average Water Usage per Acre",
            "7-Day Moving Average Pattern",
            "Meter Actual Readings (Delta)",
            "Water Usage per Acre (Delta)"
        ]
        
        graph_descriptions = [
            "This graph shows the daily average water consumption in cubic meters per acre per day. It helps identify daily irrigation patterns and water usage efficiency.",
            "This graph displays the 7-day simple moving average, smoothing out daily fluctuations to reveal longer-term trends in water usage.",
            "This graph presents the actual meter reading differences (delta) between consecutive readings, showing the total water consumed between measurements.",
            "This graph normalizes the water consumption per acre, making it easier to compare water usage across farms of different sizes."
        ]
        
        for graph_idx, img_base64 in enumerate(block['plots']):
            # Add graph title
            graph_title = doc.add_paragraph()
            graph_title.add_run(f'Graph {graph_idx + 1}: {graph_titles[graph_idx]}').bold = True
            
            # Add graph description
            doc.add_paragraph(graph_descriptions[graph_idx])
            
            # Decode and add image
            try:
                img_data = base64.b64decode(img_base64)
                img_stream = BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(6))
                doc.add_paragraph()  # Empty line after image
            except Exception as e:
                doc.add_paragraph(f"[Error loading graph: {str(e)}]")
        
        # Add page break after each meter (except last)
        if idx < len(results) - 1:
            doc.add_page_break()
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('--- End of Report ---').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer



def generate_group_analysis_report(group_type, selected_groups, group_plot_base64, group_plot2, group_data):
    """
    Generate a Word document report for group analysis
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import base64
    from datetime import datetime
    
    # Create document
    doc = Document()
    
    # Set up styles
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 139)
    
    heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 139)
    
    # Add title
    title = doc.add_paragraph('DIGI-VI Group Analysis Report', style='CustomTitle')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Add analysis information
    analysis_heading = doc.add_paragraph('Group Analysis Configuration', style='CustomHeading')
    
    # Group type information
    group_info = doc.add_paragraph()
    group_info.add_run('Study Type: ').bold = True
    group_info.add_run(f'{group_type}\n')
    
    # Selected groups
    group_info.add_run('Selected Groups: ').bold = True
    group_info.add_run(f'{", ".join(selected_groups)}\n')
    
    # Add group descriptions based on type
    doc.add_paragraph()
    desc_heading = doc.add_paragraph('Group Descriptions', style='CustomHeading')
    
    if group_type == "Remote":
        descriptions = {
            "Group-A Complied": "Treatment group with remote controllers who complied with irrigation scheduling",
            "Group-A Non-Complied": "Treatment group with remote controllers who did not comply with scheduling",
            "Group-B Complied": "Control group who complied with traditional irrigation practices",
            "Group-B Non-Complied": "Control group who did not comply with traditional practices"
        }
    elif group_type == "AWD":
        descriptions = {
            "Group-A Complied": "Treatment group practicing Alternate Wetting and Drying (AWD) with compliance",
            "Group-A Non-Complied": "Treatment group assigned AWD but with non-compliance",
            "Group-B Complied": "Control group B with compliance to assigned practices",
            "Group-B Non-Complied": "Control group B with non-compliance",
            "Group-C Complied": "Control group C with compliance to assigned practices",
            "Group-C Non-Complied": "Control group C with non-compliance"
        }
    else:  # TPR/DSR
        descriptions = {
            "TPR": "Transplanted Rice (TPR) - Traditional method of rice cultivation",
            "DSR": "Direct Seeded Rice (DSR) - Water-efficient method of rice cultivation"
        }
    
    for group in selected_groups:
        full_group_name = f"{group_type} {group}" if group_type != "TPR/DSR" else group
        if group in descriptions:
            para = doc.add_paragraph()
            para.add_run(f'{full_group_name}: ').bold = True
            para.add_run(descriptions[group])
    
    doc.add_page_break()
    
    # Add analysis results
    results_heading = doc.add_paragraph('Comparative Analysis Results', style='CustomHeading')
    
    # Add explanation
    doc.add_paragraph(
        "The following graph shows the daily average water usage (m³/acre) comparison between "
        "the selected groups over the cultivation period. This visualization helps identify "
        "water usage patterns and efficiency differences between different farming practices."
    )
    
    doc.add_paragraph()
    
    # Add the comparative graph
    try:
        img_data = base64.b64decode(group_plot_base64)
        img_stream = BytesIO(img_data)
        doc.add_picture(img_stream, width=Inches(6.5))

        img_data = base64.b64decode(group_plot2)
        img_stream = BytesIO(img_data)
        doc.add_picture(img_stream, width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph(f"[Error loading graph: {str(e)}]")
    
    # Add data summary if available
    if group_data:
        doc.add_page_break()
        summary_heading = doc.add_paragraph('Statistical Summary', style='CustomHeading')
        
        # Add summary statistics for each group
        for group_label, farms in group_data.items():
            if farms:  # Only if there are farms in this group
                para = doc.add_paragraph()
                para.add_run(f'{group_label}:').bold = True
                para.add_run(f'\n  • Number of farms: {len(farms)}')
                para.add_run(f'\n  • Farm IDs: {", ".join(farms[:5])}{"..." if len(farms) > 5 else ""}')
    
    # Add footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('--- End of Group Analysis Report ---').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer

import plotly.graph_objects as go
import pandas as pd

def get_farm_analysis_data(farm_id, merged_df, daily_df, weekly_df):
    farm_data = merged_df[merged_df["Farm ID"] == farm_id].copy()
    farm_daily_data = daily_df[daily_df["Farm ID"] == farm_id].copy()
    farm_weekly_data = weekly_df[weekly_df["Farm ID"] == farm_id].copy()

    if farm_data.empty:
        return {"error": f"No data found for Farm ID {farm_id}"}

    available_pipes = sorted(farm_data["Pipe code ID"].dropna().unique())

    # Prepare Graph 1: Daily Water Levels
    fig1 = go.Figure()
    
    # Define different marker styles for each pipe to avoid overlap confusion
    marker_symbols = ['circle', 'square', 'diamond', 'cross', 'x', 'triangle-up', 'triangle-down', 'star']
    marker_colors = px.colors.qualitative.Set1
    
    for i, pipe in enumerate(available_pipes):
        pipe_data = farm_data[farm_data["Pipe code ID"] == pipe].copy()
        pipe_data = pipe_data.sort_values("Days from TPR")
        fig1.add_trace(go.Scatter(
            x=pipe_data["Days from TPR"],
            y=pipe_data["Water Level (mm)"],
            mode='markers',
            name=f"Pipe {pipe}",
            marker=dict(
                size=8,
                symbol=marker_symbols[i % len(marker_symbols)],
                color=marker_colors[i % len(marker_colors)],
                line=dict(width=1, color='DarkSlateGrey')
            ),
            hovertemplate="<b>Pipe %{fullData.name}</b><br>" +
                         "Days from TPR: %{x}<br>" +
                         "Water Level: %{y} mm<br>" +
                         "<extra></extra>"
        ))

    if not farm_daily_data.empty:
        farm_daily_sorted = farm_daily_data.sort_values("Days from TPR")
        fig1.add_trace(go.Scatter(
            x=farm_daily_sorted["Days from TPR"],
            y=farm_daily_sorted["Water Level (mm)"],
            mode='lines+markers',
            name="Farm Average",
            line=dict(width=3, color='black'),
            marker=dict(size=6, color='black'),
            hovertemplate="<b>Farm Average</b><br>" +
                         "Days from TPR: %{x}<br>" +
                         "Water Level: %{y} mm<br>" +
                         "<extra></extra>"
        ))

    fig1.update_layout(
        title=f"Daily Water Levels - Farm {farm_id}",
        xaxis_title="Days from Transplanting (TPR)",
        yaxis_title="PVC Water Level (mm)",
        width=1200,
        height=700,
        hovermode='closest',
        legend=dict(
            orientation="v",
            yanchor="top",
            y=1,
            xanchor="left",
            x=1.02,
            # Enable multi-select by clicking legend items
            itemclick="toggleothers",  # Click to show/hide individual traces
            itemdoubleclick="toggle"   # Double-click to isolate/show all
        ),
        updatemenus=[
            dict(
                type="buttons",
                direction="down",
                showactive=True,
                x=0.02,
                y=1.15,
                buttons=[
                    dict(label="Show All Pipes", method="restyle", args=["visible", [True] * len(available_pipes)]),
                    dict(label="Hide All Pipes", method="restyle", args=["visible", ["legendonly"] * len(available_pipes)]),
                ]
            )
        ],
        annotations=[
            dict(
                text="<b>Tip:</b> Click legend items to show/hide pipes. Use buttons above for quick actions.",
                showarrow=False,
                x=0.02,
                y=1.08,
                xref="paper",
                yref="paper",
                align="left",
                font=dict(size=10, color="gray")
            )
        ]
    )

    # Prepare Graph 2: Weekly Water Levels
    fig2 = go.Figure()
    
    for i, pipe in enumerate(available_pipes):
        pipe_data = farm_data[farm_data["Pipe code ID"] == pipe].copy()
        pipe_weekly = pipe_data.groupby("Week from TPR")["Water Level (mm)"].mean().reset_index()
        fig2.add_trace(go.Scatter(
            x=pipe_weekly["Week from TPR"],
            y=pipe_weekly["Water Level (mm)"],
            mode='markers',
            name=f"Pipe {pipe}",
            marker=dict(
                size=10,
                symbol=marker_symbols[i % len(marker_symbols)],
                color=marker_colors[i % len(marker_colors)],
                line=dict(width=1, color='DarkSlateGrey')
            ),
            hovertemplate="<b>Pipe %{fullData.name}</b><br>" +
                         "Week from TPR: %{x}<br>" +
                         "Water Level: %{y} mm<br>" +
                         "<extra></extra>"
        ))

    if not farm_weekly_data.empty:
        farm_weekly_sorted = farm_weekly_data.sort_values("Week from TPR")
        fig2.add_trace(go.Scatter(
            x=farm_weekly_sorted["Week from TPR"],
            y=farm_weekly_sorted["Water Level (mm)"],
            mode='lines+markers',
            name="Farm Weekly Average",
            line=dict(width=3, color='black'),
            marker=dict(size=8, color='black'),
            hovertemplate="<b>Farm Weekly Average</b><br>" +
                         "Week from TPR: %{x}<br>" +
                         "Water Level: %{y} mm<br>" +
                         "<extra></extra>"
        ))

    fig2.update_layout(
        title=f"Weekly Water Level Trends - Farm {farm_id}",
        xaxis_title="Weeks from Transplanting (TPR)",
        yaxis_title="PVC Water Level (mm)",
        width=1200,
        height=700,
        hovermode='closest',
        legend=dict(
            orientation="v",
            yanchor="top",
            y=1,
            xanchor="left",
            x=1.02,
            # Enable multi-select by clicking legend items
            itemclick="toggleothers",  # Click to show/hide individual traces
            itemdoubleclick="toggle"   # Double-click to isolate/show all
        ),
        updatemenus=[
            dict(
                type="buttons",
                direction="down",
                showactive=True,
                x=0.02,
                y=1.15,
                buttons=[
                    dict(label="Show All Pipes", method="restyle", args=["visible", [True] * (len(available_pipes) + 1)]),
                    dict(label="Hide All Pipes", method="restyle", args=["visible", ["legendonly"] * (len(available_pipes) + 1)]),
                ]
            )
        ],
        annotations=[
            dict(
                text="<b>Tip:</b> Click legend items to show/hide pipes. Use buttons above for quick actions.",
                showarrow=False,
                x=0.02,
                y=1.08,
                xref="paper",
                yref="paper",
                align="left",
                font=dict(size=10, color="gray")
            )
        ]
    )

    # Create Table: Pivoted Water Level by Pipe per Day
    pivot_table = farm_data.pivot_table(
        index="Days from TPR",
        columns="Pipe code ID",
        values="Water Level (mm)",
        aggfunc='mean'
    ).reset_index()

    if not farm_daily_data.empty:
        farm_avg = farm_daily_data.set_index("Days from TPR")["Water Level (mm)"]
        pivot_table = pivot_table.set_index("Days from TPR")
        pivot_table["Farm Average"] = farm_avg
        pivot_table = pivot_table.reset_index()

    date_mapping = farm_data.groupby("Days from TPR")["Date"].first().reset_index()
    pivot_table = pd.merge(pivot_table, date_mapping, on="Days from TPR", how="left")

    pipe_cols = [col for col in pivot_table.columns if col not in ["Days from TPR", "Date", "Farm Average"]]
    col_order = ["Days from TPR", "Date"] + sorted(pipe_cols) + (["Farm Average"] if "Farm Average" in pivot_table.columns else [])
    pivot_table = pivot_table[col_order]

    # Summary Statistics
    summary = {
        "average": round(farm_data["Water Level (mm)"].mean(), 1),
        "std_dev": round(farm_data["Water Level (mm)"].std(), 1),
        "min": round(farm_data["Water Level (mm)"].min(), 1),
        "max": round(farm_data["Water Level (mm)"].max(), 1),
    }

    return {
        "farm_data": farm_data,
        "daily_data": farm_daily_data,
        "weekly_data": farm_weekly_data,
        "plot1": fig1.to_html(full_html=False),
        "plot2": fig2.to_html(full_html=False),
        "summary": summary,
        "pivot_table": pivot_table.to_html(index=False, classes="table table-bordered table-sm"),
        "available_pipes": available_pipes,
    }

import plotly.graph_objects as go
import plotly.express as px

def get_village_level_analysis(merged_df):
    if 'Village' not in merged_df.columns:
        return {"error": "Village column not available in merged dataset."}

    available_villages = sorted(merged_df['Village'].dropna().unique())
    filtered_df = merged_df[merged_df['Village'].isin(available_villages)]

    # Daily averages per village
    village_daily = filtered_df.groupby(['Village', 'Days from TPR']).agg({
        'Water Level (mm)': 'mean',
        'Farm ID': 'nunique'
    }).reset_index().rename(columns={'Farm ID': 'Farm Count'})

    # Plotly chart
    fig = go.Figure()
    colors = px.colors.qualitative.Set3
    marker_symbols = ['circle', 'square', 'diamond', 'cross', 'x', 'triangle-up', 'triangle-down', 'star']

    for i, village in enumerate(available_villages):
        data = village_daily[village_daily['Village'] == village].sort_values("Days from TPR")
        fig.add_trace(go.Scatter(
            x=data["Days from TPR"],
            y=data["Water Level (mm)"],
            mode="lines+markers",
            name=village,
            line=dict(color=colors[i % len(colors)], width=3),
            marker=dict(
                size=8,
                symbol=marker_symbols[i % len(marker_symbols)],
                color=colors[i % len(colors)],
                line=dict(width=1, color='DarkSlateGrey')
            ),
            hovertemplate="<b>%{fullData.name}</b><br>" +
                         "Days from TPR: %{x}<br>" +
                         "Water Level: %{y} mm<br>" +
                         "<extra></extra>"
        ))

    fig.update_layout(
        title=f"Village-level Water Level Trends Comparison ({len(available_villages)} villages)",
        xaxis_title="Days from Transplanting (TPR)",
        yaxis_title="Average Water Level (mm)",
        width=1200,
        height=700,
        hovermode="closest",
        legend=dict(
            orientation="v",
            yanchor="top",
            y=1,
            xanchor="left",
            x=1.02,
            # Enable multi-select by clicking legend items
            itemclick="toggleothers",  # Click to show/hide individual traces
            itemdoubleclick="toggle"   # Double-click to isolate/show all
        ),
        updatemenus=[
            dict(
                type="buttons",
                direction="down",
                showactive=True,
                x=0.02,
                y=1.15,
                buttons=[
                    dict(label="Show All Villages", method="restyle", args=["visible", [True] * len(available_villages)]),
                    dict(label="Hide All Villages", method="restyle", args=["visible", ["legendonly"] * len(available_villages)]),
                ]
            )
        ],
        annotations=[
            dict(
                text="<b>Tip:</b> Click legend items to show/hide villages. Use buttons above for quick actions.",
                showarrow=False,
                x=0.02,
                y=1.08,
                xref="paper",
                yref="paper",
                align="left",
                font=dict(size=10, color="gray")
            )
        ]
    )

    # Summary table
    village_summary = filtered_df.groupby('Village').agg({
        'Water Level (mm)': ['mean', 'std', 'min', 'max', 'count'],
        'Farm ID': 'nunique',
        'Days from TPR': ['min', 'max']
    }).round(2)

    village_summary.columns = [
        'Avg Water Level (mm)', 'Std Dev (mm)', 'Min Level (mm)',
        'Max Level (mm)', 'Total Readings', 'Unique Farms',
        'Min Days from TPR', 'Max Days from TPR'
    ]
    village_summary = village_summary.reset_index()

    return {
        "villages": available_villages,
        "plot": fig,
        "summary_df": village_summary
    }

import zipfile
import io

def get_remote_controllers_analysis(merged_df, kharif_df, selected_groups, mode):
    rc_cols = [
        'Kharif 25 - Remote Controllers Study - Group A - Treatment - complied (Y/N)',
        'Kharif 25 - Remote Controllers Study - Group A - Treatment - NON-complied (Y/N)',
        'Kharif 25 - Remote Controllers Study - Group B - Control - complied (Y/N)',
        'Kharif 25 - Remote Controllers Study - Group B - Control - NON-complied (Y/N)',
    ]

    study_df = pd.merge(
        merged_df,
        kharif_df[['Kharif 25 Farm ID'] + rc_cols],
        left_on='Farm ID',
        right_on='Kharif 25 Farm ID',
        how='left'
    )

    group_map = {
        'Treatment Group (A) - Complied': study_df[study_df[rc_cols[0]] == 1],
        'Treatment Group (A) - Non-Complied': study_df[study_df[rc_cols[1]] == 1],
        'Control Group (B) - Complied': study_df[study_df[rc_cols[2]] == 1],
        'Control Group (B) - Non-Complied': study_df[study_df[rc_cols[3]] == 1],
    }

    selected_data = {name: df for name, df in group_map.items() if name in selected_groups and not df.empty}

    if not selected_data:
        return {"error": "No matching Remote Controllers group data found."}

    fig = go.Figure()
    summary_data = []
    colors = px.colors.qualitative.Set1
    marker_symbols = ['circle', 'square', 'diamond', 'cross']

    if mode == "Selected Groups Comparison":
        for i, (name, df) in enumerate(selected_data.items()):
            avg = df.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
            fig.add_trace(go.Scatter(
                x=avg["Days from TPR"],
                y=avg["Water Level (mm)"],
                mode="lines+markers",
                name=name,
                line=dict(color=colors[i % len(colors)], width=3),
                marker=dict(
                    size=8,
                    symbol=marker_symbols[i % len(marker_symbols)],
                    color=colors[i % len(colors)],
                    line=dict(width=1, color='DarkSlateGrey')
                ),
                hovertemplate="<b>%{fullData.name}</b><br>" +
                             "Days from TPR: %{x}<br>" +
                             "Water Level: %{y} mm<br>" +
                             "<extra></extra>"
            ))

    elif mode == "Treatment vs Control (Selected Groups)":
        treat = pd.concat([df for name, df in selected_data.items() if "Treatment" in name])
        control = pd.concat([df for name, df in selected_data.items() if "Control" in name])

        if not treat.empty:
            avg = treat.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
            fig.add_trace(go.Scatter(
                x=avg["Days from TPR"],
                y=avg["Water Level (mm)"],
                mode="lines+markers",
                name="Treatment Combined",
                line=dict(color='blue', width=3),
                marker=dict(size=8, symbol='circle', color='blue', line=dict(width=1, color='DarkSlateGrey')),
                hovertemplate="<b>Treatment Combined</b><br>" +
                             "Days from TPR: %{x}<br>" +
                             "Water Level: %{y} mm<br>" +
                             "<extra></extra>"
            ))
        if not control.empty:
            avg = control.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
            fig.add_trace(go.Scatter(
                x=avg["Days from TPR"],
                y=avg["Water Level (mm)"],
                mode="lines+markers",
                name="Control Combined",
                line=dict(color='red', width=3),
                marker=dict(size=8, symbol='square', color='red', line=dict(width=1, color='DarkSlateGrey')),
                hovertemplate="<b>Control Combined</b><br>" +
                             "Days from TPR: %{x}<br>" +
                             "Water Level: %{y} mm<br>" +
                             "<extra></extra>"
            ))

    elif mode == "Complied vs Non-Complied (Selected Groups)":
        comp = pd.concat([df for name, df in selected_data.items() if "Complied" in name and "Non" not in name])
        noncomp = pd.concat([df for name, df in selected_data.items() if "Non-Complied" in name])

        if not comp.empty:
            avg = comp.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
            fig.add_trace(go.Scatter(
                x=avg["Days from TPR"],
                y=avg["Water Level (mm)"],
                mode="lines+markers",
                name="Complied",
                line=dict(color='green', width=3),
                marker=dict(size=8, symbol='circle', color='green', line=dict(width=1, color='DarkSlateGrey')),
                hovertemplate="<b>Complied</b><br>" +
                             "Days from TPR: %{x}<br>" +
                             "Water Level: %{y} mm<br>" +
                             "<extra></extra>"
            ))
        if not noncomp.empty:
            avg = noncomp.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
            fig.add_trace(go.Scatter(
                x=avg["Days from TPR"],
                y=avg["Water Level (mm)"],
                mode="lines+markers",
                name="Non-Complied",
                line=dict(color='orange', width=3),
                marker=dict(size=8, symbol='square', color='orange', line=dict(width=1, color='DarkSlateGrey')),
                hovertemplate="<b>Non-Complied</b><br>" +
                             "Days from TPR: %{x}<br>" +
                             "Water Level: %{y} mm<br>" +
                             "<extra></extra>"
            ))

    # Update layout with proper axis titles and dropdown
    fig.update_layout(
        title=f"Remote Controllers Analysis - {mode}",
        xaxis_title="Days from Transplanting (TPR)",
        yaxis_title="Average Water Level (mm)",
        width=1200,
        height=700,
        hovermode="closest",
        legend=dict(
            orientation="v",
            yanchor="top",
            y=1,
            xanchor="left",
            x=1.02,
            # Enable multi-select by clicking legend items
            itemclick="toggleothers",  # Click to show/hide individual traces
            itemdoubleclick="toggle"   # Double-click to isolate/show all
        ),
        updatemenus=[
            dict(
                type="buttons",
                direction="down",
                showactive=True,
                x=0.02,
                y=1.15,
                buttons=[
                    dict(label="Show All Groups", method="restyle", args=["visible", [True] * len(fig.data)]),
                    dict(label="Hide All Groups", method="restyle", args=["visible", ["legendonly"] * len(fig.data)]),
                ]
            )
        ],
        annotations=[
            dict(
                text="<b>Tip:</b> Click legend items to show/hide groups. Use buttons above for quick actions.",
                showarrow=False,
                x=0.02,
                y=1.08,
                xref="paper",
                yref="paper",
                align="left",
                font=dict(size=10, color="gray")
            )
        ]
    )

    # Summary stats - Fix the nunique() calls
    for name, df in selected_data.items():
        summary_data.append({
            'Group': name,
            'Farms': len(pd.Series(df['Farm ID']).dropna().unique()),
            'Villages': len(pd.Series(df['Village']).dropna().unique()),
            'Readings': len(df),
            'Avg Water Level': round(df['Water Level (mm)'].mean(), 1),
            'Std Dev': round(df['Water Level (mm)'].std(), 1)
        })

    summary_df = pd.DataFrame(summary_data)

    return {
        "plot": fig,
        "summary_df": summary_df,
        "selected_data": selected_data
    }

def get_awd_groups_analysis(merged_df, kharif_df, selected_groups=None, analysis_option="Selected Groups Comparison"):
    """AWD Study: Groups A, B, C Comparisons - returns fig + summary_df"""
    study_data = pd.merge(
        merged_df,
        kharif_df[[
            'Kharif 25 Farm ID',
            'Kharif 25 - AWD Study - Group A - Treatment - complied (Y/N)',
            'Kharif 25 - AWD Study - Group A - Treatment - Non-complied (Y/N)',
            'Kharif 25 - AWD Study - Group B - Complied (Y/N)',
            'Kharif 25 - AWD Study - Group B - Non-complied (Y/N)',
            'Kharif 25 - AWD Study - Group C - Complied (Y/N)',
            'Kharif 25 - AWD Study - Group C - non-complied (Y/N)'
        ]],
        left_on='Farm ID',
        right_on='Kharif 25 Farm ID',
        how='left'
    )

    groups_data = {
        'Group A (Treatment) - Complied': study_data[study_data['Kharif 25 - AWD Study - Group A - Treatment - complied (Y/N)'] == 1],
        'Group A (Treatment) - Non-Complied': study_data[study_data['Kharif 25 - AWD Study - Group A - Treatment - Non-complied (Y/N)'] == 1],
        'Group B (Training) - Complied': study_data[study_data['Kharif 25 - AWD Study - Group B - Complied (Y/N)'] == 1],
        'Group B (Training) - Non-Complied': study_data[study_data['Kharif 25 - AWD Study - Group B - Non-complied (Y/N)'] == 1],
        'Group C (Control) - Complied': study_data[study_data['Kharif 25 - AWD Study - Group C - Complied (Y/N)'] == 1],
        'Group C (Control) - Non-Complied': study_data[study_data['Kharif 25 - AWD Study - Group C - non-complied (Y/N)'] == 1],
    }

    available_groups = {k: v for k, v in groups_data.items() if not v.empty}
    if not available_groups:
        return {"error": "No AWD study data found."}

    if not selected_groups:
        selected_groups = list(available_groups.keys())[:3]

    selected_groups_data = {k: v for k, v in available_groups.items() if k in selected_groups}

    fig = go.Figure()
    colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b']
    marker_symbols = ['circle', 'square', 'diamond', 'cross', 'x', 'triangle-up']

    if analysis_option == "Complied vs Non-Complied (Selected Groups Only)":
        complied = pd.concat([v for k, v in selected_groups_data.items() if 'Complied' in k and 'Non-Complied' not in k])
        non_complied = pd.concat([v for k, v in selected_groups_data.items() if 'Non-Complied' in k])

        if not complied.empty:
            d1 = complied.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
            fig.add_trace(go.Scatter(
                x=d1["Days from TPR"], 
                y=d1["Water Level (mm)"], 
                name="Complied", 
                mode="lines+markers", 
                line=dict(color="green", width=3),
                marker=dict(size=8, symbol='circle', color='green', line=dict(width=1, color='DarkSlateGrey')),
                hovertemplate="<b>Complied</b><br>" +
                             "Days from TPR: %{x}<br>" +
                             "Water Level: %{y} mm<br>" +
                             "<extra></extra>"
            ))

        if not non_complied.empty:
            d2 = non_complied.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
            fig.add_trace(go.Scatter(
                x=d2["Days from TPR"], 
                y=d2["Water Level (mm)"], 
                name="Non-Complied", 
                mode="lines+markers", 
                line=dict(color="red", width=3),
                marker=dict(size=8, symbol='square', color='red', line=dict(width=1, color='DarkSlateGrey')),
                hovertemplate="<b>Non-Complied</b><br>" +
                             "Days from TPR: %{x}<br>" +
                             "Water Level: %{y} mm<br>" +
                             "<extra></extra>"
            ))

        fig.update_layout(
            title="Complied vs Non-Complied (AWD Groups)", 
            xaxis_title="Days from Transplanting (TPR)",
            yaxis_title="Average Water Level (mm)",
            width=1200,
            height=700,
            hovermode="closest",
            legend=dict(orientation="v", yanchor="top", y=1, xanchor="left", x=1.02)
        )

    else:  # Selected Groups Comparison
        for i, (group, data) in enumerate(selected_groups_data.items()):
            if not data.empty:
                avg = data.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
                fig.add_trace(go.Scatter(
                    x=avg["Days from TPR"], 
                    y=avg["Water Level (mm)"], 
                    name=group, 
                    mode="lines+markers", 
                    line=dict(color=colors[i % len(colors)], width=3),
                    marker=dict(
                        size=8, 
                        symbol=marker_symbols[i % len(marker_symbols)], 
                        color=colors[i % len(colors)],
                        line=dict(width=1, color='DarkSlateGrey')
                    ),
                    hovertemplate="<b>%{fullData.name}</b><br>" +
                                 "Days from TPR: %{x}<br>" +
                                 "Water Level: %{y} mm<br>" +
                                 "<extra></extra>"
                ))

        fig.update_layout(
            title=f"AWD Group Comparison ({len(selected_groups)} groups)", 
            xaxis_title="Days from Transplanting (TPR)",
            yaxis_title="Average Water Level (mm)",
            width=1200,
            height=700,
            hovermode="closest",
            legend=dict(
                orientation="v", 
                yanchor="top", 
                y=1, 
                xanchor="left", 
                x=1.02,
                # Enable multi-select by clicking legend items
                itemclick="toggleothers",  # Click to show/hide individual traces
                itemdoubleclick="toggle"   # Double-click to isolate/show all
            ),
            updatemenus=[
                dict(
                    type="buttons",
                    direction="down",
                    showactive=True,
                    x=0.02,
                    y=1.15,
                    buttons=[
                        dict(label="Show All Groups", method="restyle", args=["visible", [True] * len(selected_groups_data)]),
                        dict(label="Hide All Groups", method="restyle", args=["visible", ["legendonly"] * len(selected_groups_data)]),
                    ]
                )
            ],
            annotations=[
                dict(
                    text="<b>Tip:</b> Click legend items to show/hide groups. Use buttons above for quick actions.",
                    showarrow=False,
                    x=0.02,
                    y=1.08,
                    xref="paper",
                    yref="paper",
                    align="left",
                    font=dict(size=10, color="gray")
                )
            ]
        )

    # Summary
    summary_rows = []
    for group, data in selected_groups_data.items():
        summary_rows.append({
            "Group": group,
            "Farms": len(pd.Series(data["Farm ID"]).dropna().unique()),
            "Villages": len(pd.Series(data["Village"]).dropna().unique()),
            "Readings": len(data),
            "Avg Level (mm)": round(data["Water Level (mm)"].mean(), 1),
            "Std Dev": round(data["Water Level (mm)"].std(), 1),
        })

    summary_df = pd.DataFrame(summary_rows)

    return {
        "plot": fig,
        "summary_df": summary_df
    }

def get_dsr_tpr_analysis(merged_df, kharif_df):
    """DSR vs TPR Group Comparisons - returns fig + summary_df"""
    study_data = pd.merge(
        merged_df,
        kharif_df[[
            'Kharif 25 Farm ID',
            'Kharif 25 - DSR farm Study (Y/N)',
            'Kharif 25 - TPR Group Study (Y/N)'
        ]],
        left_on='Farm ID',
        right_on='Kharif 25 Farm ID',
        how='left'
    )

    dsr = study_data[study_data['Kharif 25 - DSR farm Study (Y/N)'] == 1]
    tpr = study_data[study_data['Kharif 25 - TPR Group Study (Y/N)'] == 1]

    fig = go.Figure()
    if not dsr.empty:
        d1 = dsr.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
        fig.add_trace(go.Scatter(x=d1["Days from TPR"], y=d1["Water Level (mm)"], name="DSR", mode="lines+markers", line=dict(color="green")))

    if not tpr.empty:
        d2 = tpr.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
        fig.add_trace(go.Scatter(x=d2["Days from TPR"], y=d2["Water Level (mm)"], name="TPR", mode="lines+markers", line=dict(color="blue")))

    fig.update_layout(title="Farming Methods: DSR vs TPR", height=500)

    summary = []
    if not dsr.empty:
        summary.append({
            "Group": "DSR",
            "Farms": len(pd.Series(dsr["Farm ID"]).dropna().unique()),
            "Villages": len(pd.Series(dsr["Village"]).dropna().unique()),
            "Readings": len(dsr),
            "Avg Level (mm)": round(dsr["Water Level (mm)"].mean(), 1),
            "Std Dev": round(dsr["Water Level (mm)"].std(), 1),
        })

    if not tpr.empty:
        summary.append({
            "Group": "TPR",
            "Farms": len(pd.Series(tpr["Farm ID"]).dropna().unique()),
            "Villages": len(pd.Series(tpr["Village"]).dropna().unique()),
            "Readings": len(tpr),
            "Avg Level (mm)": round(tpr["Water Level (mm)"].mean(), 1),
            "Std Dev": round(tpr["Water Level (mm)"].std(), 1),
        })

    return {
        "plot": fig,
        "summary_df": pd.DataFrame(summary)
    }


def compute_compliance_analysis(merged_df, kharif_df):
    """Returns compliance DataFrame and metrics for each study group."""

    study_data = pd.merge(
        merged_df,
        kharif_df[[
            'Kharif 25 Farm ID',
            'Kharif 25 - Remote Controllers Study - Group A - Treatment (Y/N)',
            'Kharif 25 - Remote Controllers Study - Group A - Treatment - complied (Y/N)',
            'Kharif 25 - Remote Controllers Study - Group B - Control Group (Y/N)',
            'Kharif 25 - Remote Controllers Study - Group B - Control - complied (Y/N)',
            'Kharif 25 - AWD Study - Group A - Treatment (Y/N)',
            'Kharif 25 - AWD Study - Group A - Treatment - complied (Y/N)',
            'Kharif 25 - AWD Study - Group B -training only (Y/N)',
            'Kharif 25 - AWD Study - Group B - Complied (Y/N)',
            'Kharif 25 - AWD Study - Group C - Control (Y/N)',
            'Kharif 25 - AWD Study - Group C - Complied (Y/N)'
        ]],
        left_on='Farm ID',
        right_on='Kharif 25 Farm ID',
        how='left'
    )

    studies = [
        {'study': 'RC - Treatment (A)', 'total_col': 'Kharif 25 - Remote Controllers Study - Group A - Treatment (Y/N)',
         'complied_col': 'Kharif 25 - Remote Controllers Study - Group A - Treatment - complied (Y/N)'},
        {'study': 'RC - Control (B)', 'total_col': 'Kharif 25 - Remote Controllers Study - Group B - Control Group (Y/N)',
         'complied_col': 'Kharif 25 - Remote Controllers Study - Group B - Control - complied (Y/N)'},
        {'study': 'AWD - Group A (Treatment)', 'total_col': 'Kharif 25 - AWD Study - Group A - Treatment (Y/N)',
         'complied_col': 'Kharif 25 - AWD Study - Group A - Treatment - complied (Y/N)'},
        {'study': 'AWD - Group B (Training)', 'total_col': 'Kharif 25 - AWD Study - Group B -training only (Y/N)',
         'complied_col': 'Kharif 25 - AWD Study - Group B - Complied (Y/N)'},
        {'study': 'AWD - Group C (Control)', 'total_col': 'Kharif 25 - AWD Study - Group C - Control (Y/N)',
         'complied_col': 'Kharif 25 - AWD Study - Group C - Complied (Y/N)'},
    ]

    compliance_data = []
    for s in studies:
        total = len(pd.Series(study_data[study_data[s['total_col']] == 1]['Farm ID']).dropna().unique())
        complied = len(pd.Series(study_data[study_data[s['complied_col']] == 1]['Farm ID']).dropna().unique())
        if total > 0:
            compliance_data.append({
                "Study Group": s['study'],
                "Total Farms": total,
                "Complied Farms": complied,
                "Non-Complied Farms": total - complied,
                "Compliance Rate (%)": round(complied / total * 100, 1)
            })

    df = pd.DataFrame(compliance_data)

    summary = {
        "overall_compliance": round(df["Complied Farms"].sum() / df["Total Farms"].sum() * 100, 1) if not df.empty else 0,
        "total_participants": df["Total Farms"].sum(),
        "total_complied": df["Complied Farms"].sum(),
        "best_group": df.loc[df["Compliance Rate (%)"].idxmax()].to_dict() if not df.empty else {},
        "worst_group": df.loc[df["Compliance Rate (%)"].idxmin()].to_dict() if not df.empty else {}
    }

    return df, summary

def summarize_custom_groups(merged_df, custom_groups: dict):
    """Return stats and daily average trends for multiple custom groups."""
    summary = []
    daily_trends = []

    for group_name, df in custom_groups.items():
        # Skip invalid groups
        if df.empty or not {'Farm ID', 'Village', 'Water Level (mm)', 'Days from TPR', 'Date'}.issubset(df.columns):
            print(f"[WARN] Skipping invalid custom group: {group_name}")
            continue

        # Convert date if needed
        if not pd.api.types.is_datetime64_any_dtype(df["Date"]):
            try:
                df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
            except Exception as e:
                print(f"[ERROR] Date parsing failed for {group_name}:", e)
                continue

        # Skip if date parsing fails completely
        if df["Date"].isnull().all():
            print(f"[WARN] All dates are NaT in group: {group_name}")
            continue

        # Daily average
        daily_avg = df.groupby("Days from TPR")["Water Level (mm)"].mean().reset_index()
        daily_avg["Group"] = group_name
        daily_trends.append(daily_avg)

        # Summary row
        summary.append({
            "Group": group_name,
            "Farms": df["Farm ID"].nunique(),
            "Villages": df["Village"].nunique(),
            "Total Readings": len(df),
            "Avg Water Level (mm)": round(df["Water Level (mm)"].mean(), 1),
            "Std Dev (mm)": round(df["Water Level (mm)"].std(), 1),
            "Date Range": f"{df['Date'].min().date()} to {df['Date'].max().date()}"
        })

    summary_df = pd.DataFrame(summary)
    trends_df = pd.concat(daily_trends, ignore_index=True) if daily_trends else pd.DataFrame()

 #Create Plotly plot
    if not trends_df.empty:
        fig = px.line(
            trends_df,
            x="Days from TPR",
            y="Water Level (mm)",
            color="Group",
            title="📈 Daily Average Water Level by Group",
            markers=True
        )
        fig.update_layout(margin=dict(l=30, r=30, t=50, b=30), height=450, paper_bgcolor="white", plot_bgcolor="white")
        fig.update_xaxes(showgrid=True, gridcolor='#e0e0e0')
        fig.update_yaxes(showgrid=True, gridcolor='#e0e0e0')
        chart_html = fig.to_html(full_html=False)
    else:
        chart_html = None

    return summary_df, trends_df, chart_html



def render_comparative_analysis(merged_df, kharif_df, custom_groups=None):
    """Django-compatible comparative analysis logic for all group types."""
    results = {}

    # 1. Village-level Aggregation
    results['village'] = get_village_level_analysis(merged_df)

    # 2. Remote Controllers: Treatment vs Control (Complied Groups)
    rc_groups = [
        "Treatment Group (A) - Complied",
        "Control Group (B) - Complied"
    ]
    rc_result = get_remote_controllers_analysis(
        merged_df, kharif_df, rc_groups, mode="Treatment vs Control (Selected Groups)"
    )
    results['rc'] = rc_result

    # 3. AWD Study: All Complied & Non-Complied
    awd_groups = [
        "Group A (Treatment) - Complied", "Group A (Treatment) - Non-Complied",
        "Group B (Training) - Complied", "Group B (Training) - Non-Complied",
        "Group C (Control) - Complied", "Group C (Control) - Non-Complied"
    ]
    awd_result = get_awd_groups_analysis(
        merged_df, kharif_df, selected_groups=awd_groups
    )
    results['awd'] = awd_result

    # 4. DSR vs TPR Group Comparisons
    dsr_tpr_result = get_dsr_tpr_analysis(merged_df, kharif_df)
    results['dsr_tpr'] = dsr_tpr_result

    # 5. Comprehensive Compliance Analysis
    compliance_df, compliance_summary = compute_compliance_analysis(merged_df, kharif_df)
    results['compliance'] = {
        "df": compliance_df,
        "summary": compliance_summary
    }

    # 6. Custom Village/Farm Selection
    if custom_groups:
        custom_summary_df, daily_trend_df, chart_html = summarize_custom_groups(merged_df, custom_groups)

        if custom_summary_df.empty or daily_trend_df.empty:
            print("[WARN] All custom groups skipped or invalid.")
            results['custom'] = {
                "summary_df": pd.DataFrame(),
                "trend_df": pd.DataFrame(),
                "chart_html": None
            }
        else:
            results['custom'] = {
                "summary_df": custom_summary_df,
                "trend_df": daily_trend_df,
                "chart_html": chart_html
            }
    else:
        results['custom'] = {
            "summary_df": pd.DataFrame(),
            "trend_df": pd.DataFrame(),
            "chart_html": None
        }

    return results



def prepare_comprehensive_downloads(merged_df, kharif_df, farm_daily_avg, weekly_avg):
    downloadables = {}

    # 1. Main datasets
    downloadables["01_complete_merged_data.csv"] = merged_df.to_csv(index=False)
    downloadables["02_raw_kharif_data.csv"] = kharif_df.to_csv(index=False)

    # 2. Processed summaries
    if not farm_daily_avg.empty:
        downloadables["03_farm_daily_averages.csv"] = farm_daily_avg.to_csv(index=False)
    
    if not weekly_avg.empty:
        downloadables["04_weekly_averages.csv"] = weekly_avg.to_csv(index=False)

    # 3. Village summary
    if 'Village' in merged_df.columns:
        village_summary = merged_df.groupby('Village').agg({
            'Water Level (mm)': ['mean', 'std', 'count'],
            'Farm ID': 'nunique'
        }).round(2)
        village_summary.columns = ['Avg Water Level', 'Std Dev', 'Total Readings', 'Unique Farms']
        downloadables["20_village_summary.csv"] = village_summary.to_csv()

    # 4. Farm summary
    farm_summary = merged_df.groupby('Farm ID').agg({
        'Water Level (mm)': ['mean', 'std', 'count'],
        'Days from TPR': ['min', 'max']
    }).round(2)
    farm_summary.columns = ['Avg Water Level', 'Std Dev', 'Total Readings', 'Min Days', 'Max Days']
    downloadables["21_farm_summary.csv"] = farm_summary.to_csv()

    # 5. Study group datasets (individual files)
    groups = create_study_group_datasets(merged_df, kharif_df)
    for i, (group_name, group_data) in enumerate(groups.items(), 5):
        filename = f"{i:02d}_{group_name.lower().replace(' ', '_').replace('(', '').replace(')', '')}.csv"
        downloadables[filename] = group_data.to_csv(index=False)

    # 6. Study Group Summaries
    if 'Remote Controller' in merged_df.columns:
        rc_summary = merged_df.groupby('Remote Controller').agg({
            'Water Level (mm)': ['mean', 'std', 'count']
        }).round(2)
        rc_summary.columns = ['Avg Water Level', 'Std Dev', 'Total Readings']
        downloadables["30_remote_controller_summary.csv"] = rc_summary.to_csv()

    if 'AWD Study' in merged_df.columns:
        awd_summary = merged_df.groupby('AWD Study').agg({
            'Water Level (mm)': ['mean', 'std', 'count']
        }).round(2)
        awd_summary.columns = ['Avg Water Level', 'Std Dev', 'Total Readings']
        downloadables["31_awd_study_summary.csv"] = awd_summary.to_csv()

    if 'Farming Method' in merged_df.columns:
        dsr_summary = merged_df.groupby('Farming Method').agg({
            'Water Level (mm)': ['mean', 'std', 'count']
        }).round(2)
        dsr_summary.columns = ['Avg Water Level', 'Std Dev', 'Total Readings']
        downloadables["32_dsr_vs_tpr_summary.csv"] = dsr_summary.to_csv()

    
    return downloadables

# --- Validation & Normalization ---

def normalize_text(text):
    if pd.isna(text) or text is None:
        return ""
    text = str(text).strip().lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s]', '', text)
    return text

def fuzzy_match_villages(village_list, threshold=85):
    """Group similar village names using simple similarity"""
    if not village_list:
        return []
    
    normalized_villages = {}
    for village in village_list:
        if pd.isna(village):
            continue
        
        norm_village = normalize_text(village)
        if not norm_village:
            continue
            
        # Check for exact matches first
        found_match = False
        for existing_norm, original in normalized_villages.items():
            if norm_village == existing_norm:
                found_match = True
                break
            # Simple similarity check
            if len(norm_village) > 3 and len(existing_norm) > 3:
                if norm_village in existing_norm or existing_norm in norm_village:
                    found_match = True
                    break
        
        if not found_match:
            normalized_villages[norm_village] = str(village).strip()
    
    return sorted(list(normalized_villages.values()))

def convert_binary_column(series):
    """Convert Y/N, Yes/No, 1/0 to consistent 1/0 format"""
    if series.dtype == 'object':
        # Handle string values
        return series.apply(lambda x: 1 if str(x).upper().strip() in ['Y', 'YES', '1', 'TRUE'] else 0)
    else:
        # Handle numeric values
        return series.apply(lambda x: 1 if pd.notna(x) and (x == 1 or x == '1') else 0)

def load_and_validate_data(kharif_file, water_file):
    required_kharif_cols = [
        "Kharif 25 Farm ID", "Kharif 25 Village",
        "Kharif 25 Paddy transplanting date (TPR)"
    ]
    required_water_cols = [
        "Date", "Farm ID", "Pipe code ID",
        "Measure water level inside the PVC pipe - millimeter mm"
    ]

    kharif_df = None
    water_df = None
    missing_kharif = []
    missing_water = []

    if kharif_file is not None:
        kharif_df = pd.read_excel(kharif_file)
        missing_kharif = [col for col in required_kharif_cols if col not in kharif_df.columns]

    if water_file is not None:
        water_df = pd.read_excel(water_file)
        missing_water = [col for col in required_water_cols if col not in water_df.columns]

    return kharif_df, water_df, missing_kharif, missing_water


def clean_and_process_data(kharif_df, water_df):
    kharif = kharif_df.copy()
    water = water_df.copy()

    # Normalize Farm IDs
    kharif['Kharif 25 Farm ID'] = kharif['Kharif 25 Farm ID'].astype(str).str.strip().str.upper()
    water['Farm ID'] = water['Farm ID'].astype(str).str.strip().str.upper()

    # Normalize Village names
    kharif['Kharif 25 Village'] = kharif['Kharif 25 Village'].astype(str).str.strip()
    if 'Village name' in water.columns:
        water['Village name'] = water['Village name'].astype(str).str.strip()

    kharif['Village_Normalized'] = kharif['Kharif 25 Village'].apply(normalize_text)
    if 'Village name' in water.columns:
        water['Village_Normalized'] = water['Village name'].apply(normalize_text)

    # Dates - Use actual TPR date if available, otherwise default to July 2, 2025
    kharif['Kharif 25 Paddy transplanting date (TPR)'] = pd.to_datetime(
        kharif['Kharif 25 Paddy transplanting date (TPR)'], errors='coerce'
    )
    kharif['Kharif 25 Paddy transplanting date (TPR)'].fillna(pd.Timestamp('2025-07-02'), inplace=True)

    water['Date'] = pd.to_datetime(water['Date'], errors='coerce')
    water = water.dropna(subset=['Date'])

    # Water level as numeric
    water['Water_Level_Numeric'] = pd.to_numeric(
        water['Measure water level inside the PVC pipe - millimeter mm'], errors='coerce'
    )
    water = water.dropna(subset=['Water_Level_Numeric'])

    # Convert binary columns
    binary_cols = [col for col in kharif.columns if '(Y/N)' in col]
    for col in binary_cols:
        kharif[col] = convert_binary_column(kharif[col])

    return kharif, water

# --- Merging Datasets ---
def create_merged_dataset(kharif_df, water_df):
    # Start with required columns
    essential_cols = [
        "Kharif 25 Farm ID", "Kharif 25 Village", "Kharif 25 Paddy transplanting date (TPR)"
    ]

    # Dynamically add study group columns if they exist
    optional_cols = []
    for col in ["Remote Controller", "AWD Study", "Farming Method"]:
        if col in kharif_df.columns:
            optional_cols.append(col)

    pvc_cols = [col for col in kharif_df.columns if "PVC Pipe code" in col]

    kharif_cols = essential_cols + optional_cols + pvc_cols
    kharif_subset = kharif_df[kharif_cols].copy()

    level_col = "Water_Level_Numeric"
    water_subset = water_df[["Date", "Farm ID", "Village name", "Pipe code ID", level_col]].copy()

    merged = pd.merge(water_subset, kharif_subset, how="inner", left_on="Farm ID", right_on="Kharif 25 Farm ID")

    merged.rename(columns={
        level_col: "Water Level (mm)",
        "Kharif 25 Paddy transplanting date (TPR)": "TPR Date",
        "Kharif 25 Village": "Village"
    }, inplace=True)

    # Ensure both Date and TPR Date are properly converted to datetime
    merged["Date"] = pd.to_datetime(merged["Date"], errors='coerce')
    merged["TPR Date"] = pd.to_datetime(merged["TPR Date"], errors='coerce')
    
    # Calculate days from TPR, handling NaT values
    merged["Days from TPR"] = (merged["Date"] - merged["TPR Date"]).dt.days.fillna(0).astype(int)
    merged["Week from TPR"] = (merged["Days from TPR"] / 7).astype(int)

    # Farm daily average
    farm_avg = merged.groupby(["Farm ID", "Date"]).agg({
        "Water Level (mm)": "mean", "TPR Date": "first", "Village": "first"
    }).reset_index()
    
    # Ensure datetime conversion for farm_avg as well
    farm_avg["Date"] = pd.to_datetime(farm_avg["Date"], errors='coerce')
    farm_avg["TPR Date"] = pd.to_datetime(farm_avg["TPR Date"], errors='coerce')
    farm_avg["Days from TPR"] = (farm_avg["Date"] - farm_avg["TPR Date"]).dt.days.fillna(0).astype(int)
    farm_avg["Week from TPR"] = (farm_avg["Days from TPR"] / 7).astype(int)

    # Weekly average
    weekly_avg = merged.groupby(["Farm ID", "Week from TPR"]).agg({
        "Water Level (mm)": "mean", "Village": "first"
    }).reset_index()

    return merged, farm_avg, weekly_avg


def apply_comprehensive_filters(kharif_df, water_df, filters):
    if kharif_df is None or water_df is None or filters is None:
        print("Early return: input was None")
        return [], []

    filtered_kharif = kharif_df.copy()
    filtered_water = water_df.copy()

    if 'date_range' in filters and filters['date_range']:
        try:
            start_date, end_date = filters['date_range']
            filtered_water = filtered_water[
                (filtered_water['Date'] >= pd.Timestamp(start_date)) &
                (filtered_water['Date'] <= pd.Timestamp(end_date))
            ]
        except Exception as e:
            print("Error in date_range unpacking:", e)

    if filters.get('villages'):
        selected = [normalize_text(v) for v in filters['villages']]
        if 'Village_Normalized' in filtered_kharif.columns:
            filtered_kharif = filtered_kharif[filtered_kharif['Village_Normalized'].isin(selected)]
        if 'Village_Normalized' in filtered_water.columns:
            filtered_water = filtered_water[filtered_water['Village_Normalized'].isin(selected)]

    for group_key in ['remote_controllers', 'awd_study', 'farming_method']:
        selected_group = filters.get(group_key, 'All')
        if selected_group != "All" and selected_group in filters.get('available_groups', {}):
            col = filters['available_groups'][selected_group]['column']
            filtered_kharif = filtered_kharif[filtered_kharif[col] == 1]

    if filters.get('min_readings', 1) > 1:
        counts = filtered_water['Farm ID'].value_counts()
        valid_farms = counts[counts >= filters['min_readings']].index
        filtered_water = filtered_water[filtered_water['Farm ID'].isin(valid_farms)]

    if filters.get('remove_outliers', False) and 'Water_Level_Numeric' in filtered_water.columns:
        Q1 = filtered_water['Water_Level_Numeric'].quantile(0.25)
        Q3 = filtered_water['Water_Level_Numeric'].quantile(0.75)
        IQR = Q3 - Q1
        lower = Q1 - 1.5 * IQR
        upper = Q3 + 1.5 * IQR
        filtered_water = filtered_water[
            (filtered_water['Water_Level_Numeric'] >= lower) &
            (filtered_water['Water_Level_Numeric'] <= upper)
        ]

    return filtered_kharif, filtered_water

# --- Study Groups ---

def create_study_group_datasets(merged_df, kharif_df):
    """Return a dictionary of study group names and their corresponding DataFrames (always includes all groups, even if empty)."""
    study_data = pd.merge(
        merged_df,
        kharif_df[[
            'Kharif 25 Farm ID',
            'Kharif 25 - Remote Controllers Study - Group A - Treatment (Y/N)',
            'Kharif 25 - Remote Controllers Study - Group A - Treatment - complied (Y/N)',
            'Kharif 25 - Remote Controllers Study - Group B - Control Group (Y/N)',
            'Kharif 25 - Remote Controllers Study - Group B - Control - complied (Y/N)',
            'Kharif 25 - AWD Study - Group A - Treatment (Y/N)',
            'Kharif 25 - AWD Study - Group A - Treatment - complied (Y/N)',
            'Kharif 25 - AWD Study - Group B -training only (Y/N)',
            'Kharif 25 - AWD Study - Group B - Complied (Y/N)',
            'Kharif 25 - AWD Study - Group C - Control (Y/N)',
            'Kharif 25 - AWD Study - Group C - Complied (Y/N)',
            'Kharif 25 - TPR Group Study (Y/N)',
            'Kharif 25 - DSR farm Study (Y/N)'
        ]],
        left_on='Farm ID',
        right_on='Kharif 25 Farm ID',
        how='left'
    )

    study_groups = {}
    # Always include all groups, even if empty
    study_groups['Remote Controllers Treatment'] = study_data[study_data['Kharif 25 - Remote Controllers Study - Group A - Treatment - complied (Y/N)'] == 1]
    study_groups['Remote Controllers Control'] = study_data[study_data['Kharif 25 - Remote Controllers Study - Group B - Control - complied (Y/N)'] == 1]
    study_groups['AWD Group A Treatment'] = study_data[study_data['Kharif 25 - AWD Study - Group A - Treatment - complied (Y/N)'] == 1]
    study_groups['AWD Group B Training'] = study_data[study_data['Kharif 25 - AWD Study - Group B - Complied (Y/N)'] == 1]
    study_groups['AWD Group C Control'] = study_data[study_data['Kharif 25 - AWD Study - Group C - Complied (Y/N)'] == 1]
    study_groups['DSR Farms'] = study_data[study_data['Kharif 25 - DSR farm Study (Y/N)'] == 1]
    study_groups['TPR Farms'] = study_data[study_data['Kharif 25 - TPR Group Study (Y/N)'] == 1]
    return study_groups

import zipfile
def create_zip_package(merged_df, kharif_df, farm_daily_avg, weekly_avg):
    """Returns a bytes buffer of a ZIP package containing all data."""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:

        def write_csv(name, df):
            buffer = io.StringIO()
            df.to_csv(buffer, index=False)
            zip_file.writestr(name, buffer.getvalue())

        # Main datasets
        write_csv("01_complete_merged_data.csv", merged_df)
        write_csv("02_raw_kharif_data.csv", kharif_df)

        if not farm_daily_avg.empty:
            write_csv("03_farm_daily_averages.csv", farm_daily_avg)

        if not weekly_avg.empty:
            write_csv("04_weekly_averages.csv", weekly_avg)

        # Study groups
        groups = create_study_group_datasets(merged_df, kharif_df)
        for i, (group_name, group_data) in enumerate(groups.items(), 5):
            filename = f"{i:02d}_{group_name.lower().replace(' ', '_').replace('(', '').replace(')', '')}.csv"
            write_csv(filename, group_data)

        # Summary reports
        if 'Village' in merged_df.columns:
            village_summary = merged_df.groupby('Village').agg({
                'Water Level (mm)': ['mean', 'std', 'count'],
                'Farm ID': 'nunique'
            }).round(2)
            village_summary.columns = ['Avg Water Level', 'Std Dev', 'Total Readings', 'Unique Farms']
            write_csv("20_village_summary.csv", village_summary)

        farm_summary = merged_df.groupby('Farm ID').agg({
            'Water Level (mm)': ['mean', 'std', 'count'],
            'Days from TPR': ['min', 'max']
        }).round(2)
        farm_summary.columns = ['Avg Water Level', 'Std Dev', 'Total Readings', 'Min Days', 'Max Days']
        write_csv("21_farm_summary.csv", farm_summary)

        # Add README
        readme = f"""Agricultural Data Analysis Package
=================================

This package contains:

Main Datasets:
- 01_complete_merged_data.csv: All water level data merged with farm information
- 02_raw_kharif_data.csv: Original farm and study group data
- 03_farm_daily_averages.csv: Daily averages per farm
- 04_weekly_averages.csv: Weekly averages per farm

Study Groups:
- Remote Controllers Treatment/Control groups
- AWD Study Groups A, B, C
- DSR and TPR farming mewatethod groups

Summary Reports:
- 20_village_summary.csv: Statistics by village
- 21_farm_summary.csv: Statistics by farm
Generated on: """ + datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        zip_file.writestr("README.txt", readme)

    zip_buffer.seek(0)
    return zip_buffer

import io
import re

def sanitize_filename(text):
    """Replaces unsafe characters in filename"""
    return re.sub(r'[^\w\-_.]', '_', text)

def get_per_farm_downloads(merged_df, farm_daily_avg):
    """Generates downloadable CSVs for all farms with safe filenames"""
    downloads = {}
    unique_farms = merged_df['Farm ID'].unique()

    for farm_id in unique_farms:
        farm_data = merged_df[merged_df['Farm ID'] == farm_id]
        safe_farm_id = sanitize_filename(str(farm_id))  # 🧼 Clean farm ID for filename use

        # 📁 Detailed CSV
        detailed_csv = farm_data.to_csv(index=False)
        downloads[f"farm_{safe_farm_id}_detailed.csv"] = detailed_csv

        # 📊 Summary CSV (1 row per farm)
        summary_data = {
            "Farm ID": farm_id,
            "Average": farm_data["Water Level (mm)"].mean(),
            "Std Dev": farm_data["Water Level (mm)"].std(),
            "Min": farm_data["Water Level (mm)"].min(),
            "Max": farm_data["Water Level (mm)"].max(),
            "Total Readings": len(farm_data)
        }
        summary_df = pd.DataFrame([summary_data])
        downloads[f"farm_{safe_farm_id}_summary.csv"] = summary_df.to_csv(index=False)

        # 📈 Daily average CSV
        if farm_daily_avg is not None:
            daily_avg = farm_daily_avg[farm_daily_avg['Farm ID'] == farm_id]
            downloads[f"farm_{safe_farm_id}_daily_avg.csv"] = daily_avg.to_csv(index=False)

    return downloads


import plotly.graph_objects as go
from datetime import datetime

# Add these new functions to utils.py (parallel to existing matplotlib functions)

def get_2025plots_plotly(raw_df, master_df, selected_farm, meter_list, start_date_enter=None, end_date_enter=None):
    """
    Generate Plotly interactive plots for multiple meters (parallel to get_2025plots)
    Returns list of Plotly HTML strings instead of base64 images
    """
    meta = master_df['Farm details']
    farm_row = meta[meta['Kharif 25 Farm ID'] == selected_farm]
    if farm_row.empty:
        raise ValueError(f"Farm ID {selected_farm} not found in metadata.")
    farm_row = farm_row.iloc[0]

    acreage = farm_row.get('Kharif 25 Acres farm - farmer reporting') or 1
    if pd.isna(acreage) or acreage <= 0:
        acreage = 1

    if pd.notna(farm_row.get('Kharif 25 Paddy transplanting date (TPR)')):
        start_date = pd.to_datetime(farm_row['Kharif 25 Paddy transplanting date (TPR)'], dayfirst=True)
    else:
        start_date = pd.to_datetime('20/06/2025', dayfirst=True)

    if start_date_enter is not None:
        start_date = pd.to_datetime(start_date_enter)

    end_date = pd.to_datetime(datetime.now().date())
    if end_date_enter is not None:
        end_date = pd.to_datetime(end_date_enter)

    # Clean meter data
    date_col = 'Date'
    raw_df[date_col] = pd.to_datetime(raw_df[date_col], dayfirst=False)
    raw_df = raw_df.sort_values(date_col)

    plots_html = []

    for meter in meter_list:
        filled_df = get_filled(meter, raw_df, date_col, start_date, end_date, acreage)
        if filled_df.empty:
            continue

        filled_df['Day'] = (pd.to_datetime(filled_df[date_col]) - start_date).dt.days
        filled_df['7-day SMA'] = filled_df['m³ per Acre per Avg Day'].rolling(window=7, min_periods=1).mean()
        x_axis = filled_df['Date'] if start_date_enter else filled_df['Day']

        dates_for_reading = get_dates(meter, raw_df, date_col, start_date, end_date, acreage)

        marker_points = filled_df[filled_df['Date'].isin(dates_for_reading['Date'])]
        # Graph 1: m³ per Acre per Avg Day
        fig1 = go.Figure()
        fig1.add_trace(go.Scatter(
            x=x_axis,
            y=filled_df['m³ per Acre per Avg Day'],
            mode='lines',
            name='Daily m³ per Acre per day',
            line=dict(color='#3b82f6', width=2),
            hovertemplate='<b>Day %{x}</b><br>Daily Avg: %{y:.2f} m³/acre<extra></extra>'
        ))
        fig1.add_trace(go.Scatter(
            x=marker_points['Date'],
            y=marker_points['m³ per Acre per Avg Day'],
            mode='markers',
            name='Reading Dates',
            marker=dict(color='#3b82f6', size=6),
            showlegend=False,
            hovertemplate='<b>Day %{x}</b><br>Daily Avg: %{y:.2f} m³/acre<extra></extra>'
        ))

        fig1.update_layout(
            title=f'Daily Avg per Acre | Meter {meter}',
            xaxis_title='Days from transplanting',
            yaxis_title='Daily Avg m³ per Acre',
            height=450,
            margin=dict(l=60, r=60, t=80, b=60),
            hovermode='x unified',
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)'
        )
        fig1.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
        fig1.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
        plots_html.append(fig1.to_html(full_html=False, include_plotlyjs='cdn'))

        # Graph 2: Moving Averages
        fig2 = go.Figure()
        fig2.add_trace(go.Scatter(
            x=x_axis,
            y=filled_df['7-day SMA'],
            mode='lines',
            name='7-day SMA',
            line=dict(color='#10b981', width=2, dash='dash'),
            hovertemplate='<b>Day %{x}</b><br>7-day SMA: %{y:.2f} m³/acre<extra></extra>'
        ))
        fig2.add_trace(go.Scatter(
            x=marker_points['Date'],
            y=marker_points['7-day SMA'],
            mode='markers',
            name='Reading Dates',
            marker=dict(color='#3b82f6', size=6),
            showlegend=False,
            hovertemplate='<b>Day %{x}</b><br>Daily Avg: %{y:.2f} m³/acre<extra></extra>'
        ))
        fig2.update_layout(
            title=f'Moving Average | Meter {meter}',
            xaxis_title='Days from transplanting',
            yaxis_title='7-days SMA m³ per Acre',
            height=450,
            margin=dict(l=60, r=60, t=80, b=60),
            hovermode='x unified',
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)'
        )
        fig2.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
        fig2.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
        plots_html.append(fig2.to_html(full_html=False, include_plotlyjs='cdn'))

        # Graph 3: Delta Analysis
        fig3 = go.Figure()
        fig3.add_trace(go.Scatter(
            x=x_axis,
            y=filled_df['Delta m³'],
            mode='lines',
            name='Delta m³',
            line=dict(color='#8b5cf6', width=2),
            hovertemplate='<b>Day %{x}</b><br>Delta: %{y:.2f} m³<extra></extra>'
        ))
        fig3.add_trace(go.Scatter(
            x=marker_points['Date'],
            y=marker_points['Delta m³'],
            mode='markers',
            name='Reading Dates',
            marker=dict(color='#3b82f6', size=6),
            showlegend=False,
            hovertemplate='<b>Day %{x}</b><br>Daily Avg: %{y:.2f} m³/acre<extra></extra>'
        ))
        fig3.update_layout(
            title=f'Meter Actual readings | Meter {meter}',
            xaxis_title='Days from transplanting',
            yaxis_title='Delta of readings (m³)',
            height=450,
            margin=dict(l=60, r=60, t=80, b=60),
            hovermode='x unified',
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)'
        )
        fig3.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
        fig3.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
        plots_html.append(fig3.to_html(full_html=False, include_plotlyjs='cdn'))

        # Graph 4: Delta per Acre
        fig4 = go.Figure()
        fig4.add_trace(go.Scatter(
            x=x_axis,
            y=filled_df['m³ per Acre'],
            mode='lines',
            name='Delta m³/Acre',
            line=dict(color='#ef4444', width=2),
            hovertemplate='<b>Day %{x}</b><br>Per Acre: %{y:.2f} m³<extra></extra>'
        ))
        fig4.add_trace(go.Scatter(
            x=marker_points['Date'],
            y=marker_points['m³ per Acre'],
            mode='markers',
            name='Reading Dates',
            marker=dict(color='#3b82f6', size=6),
            showlegend=False,
            hovertemplate='<b>Day %{x}</b><br>Daily Avg: %{y:.2f} m³/acre<extra></extra>'
        ))
        fig4.update_layout(
            title=f'Meter readings per Acre | Meter {meter}',
            xaxis_title='Days from transplanting',
            yaxis_title='Delta of reading per Acre (m³)',
            height=450,
            margin=dict(l=60, r=60, t=80, b=60),
            hovermode='x unified',
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)'
        )
        fig4.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
        fig4.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
        plots_html.append(fig4.to_html(full_html=False, include_plotlyjs='cdn'))

    return plots_html


def get_2025plots_combined_plotly(raw_df, master_df, selected_farm, meter_list, start_date_enter=None, end_date_enter=None):
    """
    Generate combined Plotly plots for multiple meters (parallel to get_2025plots_combined)
    Returns list of Plotly HTML strings instead of base64 images
    """
    if len(meter_list) < 2:
        return []  # No combined plots for single meter

    # Get farm metadata (same logic as original)
    meta = master_df['Farm details']
    farm_row = meta[meta['Kharif 25 Farm ID'] == selected_farm]
    if farm_row.empty:
        raise ValueError(f"Farm ID {selected_farm} not found in metadata.")
    farm_row = farm_row.iloc[0]

    acreage = farm_row.get('Kharif 25 Acres farm - farmer reporting') or 1
    if pd.isna(acreage) or acreage <= 0:
        acreage = 1

    if pd.notna(farm_row.get('Kharif 25 Paddy transplanting date (TPR)')):
        start_date = pd.to_datetime(farm_row['Kharif 25 Paddy transplanting date (TPR)'], dayfirst=True)
    else:
        start_date = pd.to_datetime('20/06/2025', dayfirst=True)

    if start_date_enter is not None:
        start_date = pd.to_datetime(start_date_enter)

    end_date = pd.to_datetime(datetime.now().date())
    if end_date_enter is not None:
        end_date = pd.to_datetime(end_date_enter)

    # Clean meter data
    date_col = 'Date'
    raw_df[date_col] = pd.to_datetime(raw_df[date_col], dayfirst=False)
    raw_df = raw_df.sort_values(date_col)

    # Get data for all meters (same logic as original)
    all_meter_data = []
    for meter in meter_list:
        meter_df = raw_df[raw_df['Meter Serial Number - as shown on meter'] == meter].copy()
        if not meter_df.empty:
            meter_df = meter_df[[date_col, 'Reading in the meter - in m3']].dropna()
            meter_df[date_col] = pd.to_datetime(meter_df[date_col])
            meter_df = meter_df.sort_values(date_col)
            meter_df = meter_df.drop_duplicates(subset=[date_col], keep='last')
            meter_df['Meter'] = meter
            all_meter_data.append(meter_df)

    if not all_meter_data:
        return []

    # Combine all meter data (same logic)
    combined_df = pd.concat(all_meter_data, ignore_index=True)
    combined_grouped = combined_df.groupby(date_col).agg({
        'Reading in the meter - in m3': 'sum'
    }).reset_index()

    combined_grouped = combined_grouped.sort_values(date_col)
    combined_grouped['Days Since Previous Reading'] = combined_grouped[date_col].diff().dt.days.fillna(1).astype(int)
    combined_grouped['Delta m³'] = combined_grouped['Reading in the meter - in m3'].diff().fillna(0)

    # Normalize per acre
    combined_grouped['m³ per Acre'] = combined_grouped['Delta m³'] / acreage
    combined_grouped['m³ per Acre per Avg Day'] = combined_grouped['m³ per Acre'] / combined_grouped['Days Since Previous Reading'].replace(0, 1)

    # Fill daily data
    all_dates = pd.date_range(start=start_date, end=end_date, freq='D')
    base = pd.DataFrame({date_col: all_dates})
    filled_df = pd.merge(base, combined_grouped, on=date_col, how='left')

    # Forward fill values
    for col in ['m³ per Acre', 'm³ per Acre per Avg Day', 'Delta m³']:
        filled_df[col] = filled_df[col].fillna(0)

    # Calculate days from start
    filled_df['Day'] = (filled_df[date_col] - start_date).dt.days
    filled_df['7-day SMA'] = filled_df['m³ per Acre per Avg Day'].rolling(window=7, min_periods=1).mean()

    # Generate plots
    plots_html = []
    meters_label = " + ".join(meter_list)

    # Graph 1: Daily avg per acre
    fig1 = go.Figure()
    fig1.add_trace(go.Scatter(
        x=filled_df['Day'],
        y=filled_df['m³ per Acre per Avg Day'],
        mode='lines+markers',
        name='Combined Daily m³ per Acre per day',
        line=dict(color='#3b82f6', width=3),
        marker=dict(size=4),
        hovertemplate='<b>Day %{x}</b><br>Combined Daily Avg: %{y:.2f} m³/acre<extra></extra>'
    ))
    fig1.update_layout(
        title=f'Combined Daily Avg per Acre | Meters: {meters_label}',
        xaxis_title='Days from transplanting',
        yaxis_title='Daily Avg m³ per Acre',
        height=450,
        margin=dict(l=60, r=60, t=80, b=60),
        hovermode='x unified',
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)'
    )
    fig1.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
    fig1.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
    plots_html.append(fig1.to_html(full_html=False, include_plotlyjs='cdn'))

    # Graph 2: 7-day SMA
    fig2 = go.Figure()
    fig2.add_trace(go.Scatter(
        x=filled_df['Day'],
        y=filled_df['7-day SMA'],
        mode='lines+markers',
        name='Combined 7-day SMA',
        line=dict(color='#10b981', width=3, dash='dash'),
        marker=dict(size=4),
        hovertemplate='<b>Day %{x}</b><br>Combined 7-day SMA: %{y:.2f} m³/acre<extra></extra>'
    ))
    fig2.update_layout(
        title=f'Combined Moving Average | Meters: {meters_label}',
        xaxis_title='Days from transplanting',
        yaxis_title='7-days SMA m³ per Acre',
        height=450,
        margin=dict(l=60, r=60, t=80, b=60),
        hovermode='x unified',
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)'
    )
    fig2.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
    fig2.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
    plots_html.append(fig2.to_html(full_html=False, include_plotlyjs='cdn'))

    # Graph 3: Delta Analysis
    fig3 = go.Figure()
    fig3.add_trace(go.Scatter(
        x=filled_df['Day'],
        y=filled_df['Delta m³'],
        mode='lines+markers',
        name='Combined Delta m³',
        line=dict(color='#8b5cf6', width=3),
        marker=dict(size=6, symbol='circle'),
        hovertemplate='<b>Day %{x}</b><br>Combined Delta: %{y:.2f} m³<extra></extra>'
    ))
    fig3.update_layout(
        title=f'Combined Meter Readings | Meters: {meters_label}',
        xaxis_title='Days from transplanting',
        yaxis_title='Combined Delta of readings (m³)',
        height=450,
        margin=dict(l=60, r=60, t=80, b=60),
        hovermode='x unified',
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)'
    )
    fig3.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
    fig3.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
    plots_html.append(fig3.to_html(full_html=False, include_plotlyjs='cdn'))

    # Graph 4: Delta per Acre
    fig4 = go.Figure()
    fig4.add_trace(go.Scatter(
        x=filled_df['Day'],
        y=filled_df['m³ per Acre'],
        mode='lines+markers',
        name='Combined Delta m³/Acre',
        line=dict(color='#ef4444', width=3),
        marker=dict(size=6, symbol='x'),
        hovertemplate='<b>Day %{x}</b><br>Combined Per Acre: %{y:.2f} m³<extra></extra>'
    ))
    fig4.update_layout(
        title=f'Combined Readings per Acre | Meters: {meters_label}',
        xaxis_title='Days from transplanting',
        yaxis_title='Combined Delta per Acre (m³)',
        height=450,
        margin=dict(l=60, r=60, t=80, b=60),
        hovermode='x unified',
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)'
    )
    fig4.update_xaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
    fig4.update_yaxes(showgrid=True, gridwidth=1, gridcolor='rgba(128,128,128,0.2)')
    plots_html.append(fig4.to_html(full_html=False, include_plotlyjs='cdn'))

    return plots_html